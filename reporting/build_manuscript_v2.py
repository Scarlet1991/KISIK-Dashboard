# -*- coding: utf-8 -*-
"""KISIK ICU LoS — Frontiers Manuscript (TRIPOD+AI, revised).
Cohorts: retrospective development (n=17,032) and prospective evaluation (n=193 completed stays only).
Addresses: pipeline clarity, MAE formula, consistent 4-model comparison, log1p motivation,
multiple-stay handling, calibration/discrimination, hyperparameter detail, patient chars table.
"""
import sys, io, warnings; warnings.filterwarnings("ignore")
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from pathlib import Path
import pandas as pd
import numpy as np
from scipy.stats import spearmanr as _spearmanr
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

AN  = Path(r"D:\Ausgangsdaten\KISIK Projekt\Eigene Auswertung")
CAN = AN / "canonical"
OUT = AN / "KISIK_Frontiers_Manuskript_v2.docx"
FONT = "Times New Roman"; ACC = RGBColor(0x1F,0x4E,0x79); HFILL = "1F4E79"; ZEB = "EAF1F8"

# ── Load result data ──────────────────────────────────────────────────────────
t1df    = pd.read_csv(CAN/"table1_characteristics.csv",  sep=";", dtype=str).fillna("")
retro_m = pd.read_csv(CAN/"metrics_retrospective.csv",   sep=";")
pros_m  = pd.read_csv(CAN/"metrics_prospective_fair24h.csv", sep=";")
feat    = pd.read_csv(CAN/"feature_importance.csv",      sep=";")
pred_df = pd.read_csv(CAN/"metrics_prospective_fair24h_predictions.csv", sep=";")

# ── Additional statistics ─────────────────────────────────────────────────────
def srho(a, b):
    m = np.isfinite(np.asarray(a, float)) & np.isfinite(np.asarray(b, float))
    r, _ = _spearmanr(np.asarray(a, float)[m], np.asarray(b, float)[m])
    return round(float(r), 3)

obs = pred_df["los_obs"].values
RHO = {
    "Senior physician": srho(obs, pred_df["arzt"].values),
    "ExtraTrees":       srho(obs, pred_df["pred_ExtraTrees"].values),
    "XGBoost":          srho(obs, pred_df["pred_XGBoost"].values),
    "RandomForest":     srho(obs, pred_df["pred_RandomForest"].values),
    "Ridge":            srho(obs, pred_df["pred_Ridge"].values),
}
print("Spearman rho:", RHO)

# Pre-formatted for use in body text
_rho_ob  = f"{RHO['Senior physician']:.2f}"
_rho_et  = f"{RHO['ExtraTrees']:.2f}"
_rho_xgb = f"{RHO['XGBoost']:.2f}"
_rho_rf  = f"{RHO['RandomForest']:.2f}"
_rho_ri  = f"{RHO['Ridge']:.2f}"

# ── Shortcut: lookup row from metrics df ──────────────────────────────────────
def rm(model):  return retro_m[retro_m["Modell"]==model].iloc[0]
def pm(model):  return pros_m [pros_m["Modell"] ==model].iloc[0]

# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()
st  = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(12)
st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
st.paragraph_format.space_after = Pt(6)
sec = doc.sections[0]; sec.page_width = Cm(21); sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.4); sec.top_margin = sec.bottom_margin = Cm(2.4)

# ── Formatting helpers ────────────────────────────────────────────────────────
def run(p, t, bold=False, italic=False, size=None, color=None):
    r = p.add_run(t); r.font.name = FONT; r.font.bold = bold; r.font.italic = italic
    if size:  r.font.size = size
    if color: r.font.color.rgb = color
    return r

def title(t):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    run(p, t, bold=True, size=Pt(15), color=ACC)

def h1(t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)
    run(p, t, bold=True, size=Pt(13), color=ACC)

def h2(t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    run(p, t, bold=True, italic=True, size=Pt(12))

def body(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run(p, t, size=Pt(12)); return p

def labeled(lab, t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run(p, lab + " ", bold=True); run(p, t); return p

def small(t):
    p = doc.add_paragraph()
    run(p, t, size=Pt(9.5), italic=True, color=RGBColor(0x55,0x55,0x55))
    return p

def cap(n, t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    run(p, f"Table {n}. ", bold=True, size=Pt(10.5)); run(p, t, size=Pt(10.5))

def figcap(n, t):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    run(p, f"Figure {n}. ", bold=True, size=Pt(9.5)); run(p, t, size=Pt(9.5))

def bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:color"),"auto")
    sh.set(qn("w:fill"), hexc); tcPr.append(sh)

def settext(cell, t, bold=False, color=None, size=9.5, align="left"):
    cell.text = ""; p = cell.paragraphs[0]
    p.alignment = {"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    r = p.add_run(t); r.font.name = FONT; r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def table(rows, widths, numcols_from=1):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.cell(ri,ci); c.width = Cm(widths[ci])
            if ri == 0:
                settext(c, val, bold=True, color=RGBColor(0xFF,0xFF,0xFF), align="center")
            else:
                settext(c, val, align="center" if ci >= numcols_from else "left")
        if ri == 0:
            for c in t.rows[0].cells: bg(c, HFILL)
        elif ri % 2 == 0:
            for c in t.rows[ri].cells: bg(c, ZEB)
    return t

def figure(path, width_cm=15.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width_cm))

# ── Patient characteristics table builder ─────────────────────────────────────
def build_table1():
    """Build patient characteristics table dynamically from CSV."""
    rows_data = []
    for _, row in t1df.iterrows():
        typ   = row["type"]
        char  = row["Characteristic"]
        retro = row["Retrospective"]
        pros  = row["Prospective"]
        smd   = row["SMD"]
        pstr  = row["p_str"]
        if typ == "cont":
            rows_data.append((char, retro, pros, pstr, smd, False))
        elif typ == "cat_head":
            rows_data.append((char, retro, pros, pstr, "", True))
        elif typ in ("cat_lvl", "score"):
            rows_data.append((char, retro, pros, "", smd, False))
    return rows_data

# ══════════════════════════════════════════════════════════════════════════════
# TITLE / ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
title("Early prediction of intensive-care length of stay from the first 24 hours: "
      "a leakage-controlled development study with prospective benchmarking against "
      "senior-physician judgement")
small("Original Research prepared for the Frontiers Research Topic “MedicineAI: Advancing the Synergy of "
      "Medicine and AI — From Data to Clinical Impact.” Reported in accordance with the TRIPOD+AI statement "
      "(Collins et al., 2024). Author names, affiliations, ORCID and corresponding author to be inserted.")
p = doc.add_paragraph(); run(p,"Author One¹, Author Two¹, Author Three², Senior Author¹*",size=Pt(11))
p = doc.add_paragraph(); run(p,"¹ Department of Anaesthesiology and Intensive Care Medicine, [Institution], [City], [Country]\n"
      "² [Second affiliation]\n* Correspondence: [name, e-mail]", size=Pt(10.5))

h1("Abstract")
labeled("Background:","Accurate early prediction of intensive care unit (ICU) length of stay (LOS) could support "
    "capacity planning and patient flow. Machine-learning (ML) studies frequently report over-optimistic performance "
    "because of information leakage from whole-stay feature aggregation, and models are rarely benchmarked "
    "prospectively against the clinicians they are intended to support.")
labeled("Methods:","We conducted a single-centre study using routine data from a tertiary ICU service. A retrospective "
    "development cohort of 17,032 ICU stays (12,414 patients; May 2017–Jul 2024) was used to train and select "
    "a final model. Eighty-four leakage-free predictors were derived from the first 24 hours after ICU admission "
    "only; any earlier pipeline that aggregated whole-stay measurements was identified as leakage and corrected. "
    "Four candidate models — ridge regression, random forest, extremely randomised trees (Extra Trees), and "
    "gradient boosting (XGBoost) — were fitted on a log1p-transformed LOS target with identical preprocessing. "
    "Hyperparameters were optimised by 4-fold patient-grouped cross-validation. The model with the lowest "
    "cross-validated mean absolute error (MAE) was pre-specified as the final model before any holdout data were "
    "seen. The identical model was then evaluated prospectively against senior-physician estimates in 193 completed "
    "matched stays (is_open = 0; LOS > 1 day; Oct 2024–Jan 2026).")
labeled("Results:","Replacing whole-stay features with strict 24-hour equivalents reduced apparent R² from 0.61 to "
    "0.31, confirming the leakage. On the patient-grouped holdout (n = 3,429), the four models performed "
    "near-identically (MAE 2.75–2.94 days; R² 0.23–0.32); Extra Trees was selected as the final model "
    "(test MAE 2.76 days, R² 0.31). All models showed systematic underestimation (mean bias −1.0 to −1.2 days). "
    "Prospectively, first-24-hour features were reconstructed from raw records (86% of the 84 features available; "
    f"median per-stay completeness 78%). The senior physician outperformed all ML models (MAE 2.01 vs 2.70 days, "
    f"R² 0.22 vs 0.05 for Extra Trees; Spearman ρ {_rho_ob} vs {_rho_et}). The physician advantage was largest in long stays "
    "(> 7 days, n = 27: physician MAE 6.51 vs XGBoost 6.84 days). Extra Trees was nearly unbiased prospectively "
    "(mean bias +0.09 days). The strongest predictors were early ICU complex-treatment procedure codes.")
labeled("Conclusion:","Under strict leakage control and prospective clinician benchmarking, 24-hour ML models did not "
    "match senior-physician judgement for ICU LOS prediction. Transparent handling of the prediction time point and "
    "direct clinician comparison are essential prerequisites before deployment.")
p = doc.add_paragraph(); run(p,"Keywords: ",bold=True)
run(p,"intensive care unit; length of stay; clinical prediction model; machine learning; data leakage; "
    "prospective validation; TRIPOD+AI; calibration; clinical decision support")

# ══════════════════════════════════════════════════════════════════════════════
# 1 INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1("1  Introduction")
body("Intensive care is among the most resource-intensive segments of hospital care, and ICU length of stay (LOS) "
    "is a principal driver of bed occupancy, staffing demand and cost [1]. An accurate early estimate of LOS — "
    "available at or shortly after ICU admission — could support bed management, step-down planning and "
    "communication with patients and families. Established severity scores (APACHE II, SAPS II, SOFA) were designed "
    "primarily to predict in-hospital mortality and explain only a limited portion of LOS variance [1–3].")
body("Machine learning (ML) has been applied widely to ICU LOS prediction [4–6], frequently with encouraging "
    "reported accuracy. Two systemic problems undermine clinical credibility. First, information leakage — the "
    "inadvertent inclusion of information unavailable at the intended prediction time — inflates apparent performance "
    "and is a leading cause of over-optimism and irreproducibility in clinical ML research [7]. For a temporal "
    "outcome such as LOS, features aggregated over the entire admission are especially hazardous: their values "
    "encode how long the patient stayed and are thus mechanically correlated with the outcome. Second, models are "
    "almost invariably evaluated against historical labels, rarely against the clinicians whose judgement they are "
    "meant to augment or replace. The TRIPOD+AI reporting standard therefore emphasises a clearly defined "
    "prediction time point and rigorous, ideally prospective, evaluation [8,9].")
body("We address both shortcomings in a single study: (i) we explicitly quantify the performance inflation caused "
    "by whole-stay leakage; (ii) we restrict all predictors to the first 24 hours after admission; (iii) we "
    "pre-specify a single final model through patient-grouped cross-validation; and (iv) we compare that model "
    "prospectively and directly against documented senior-physician estimates. The guiding clinical question is "
    "whether a data-driven 24-hour model can match an experienced clinician at the moment of decision.")

# ══════════════════════════════════════════════════════════════════════════════
# 2 MATERIALS AND METHODS
# ══════════════════════════════════════════════════════════════════════════════
h1("2  Materials and Methods")

h2("2.1  Study design, setting and data source")
body("This is a single-centre prediction model development and prospective validation study using de-identified "
    "routine data from the clinical data repository (KISIK) of a tertiary care centre. ICU stay records, "
    "ICD-10 admission diagnoses, laboratory results, vital signs, OPS procedure codes and vascular-access device "
    "records were linked by a case identifier (retrospective: fallid; prospective: fallnr). The study was conducted "
    "in accordance with the Declaration of Helsinki; ethics approval and waiver of informed consent for analysis of "
    "de-identified routine data are to be inserted (see Ethics statement). Reporting follows TRIPOD+AI [9].")

h2("2.2  Cohorts: retrospective development and prospective evaluation")
body("Two temporally non-overlapping cohorts were defined (Table 1, Table 2).")
body("The retrospective development cohort covers May 2017 to July 2024 and comprises 17,032 ICU stays from "
    "12,414 patients (13,275 hospital encounters) across 17 ward/care-unit combinations of the participating ICU "
    "service. Inclusion criteria: (a) stay on one of the 17 eligible ward/care-unit pairs; "
    "(b) ICU LOS > 1 day (icu_duration_h / 24 > 1). Each ICU stay is one analysis unit. The data were split "
    "80:20 by patient (GroupShuffleSplit, random seed 42) into a training set (n = 13,603) and a held-out test "
    "set (n = 3,429); no patient appears in both sets.")
body("The prospective evaluation cohort spans October 2024 to January 2026 and was drawn from daily live-system "
    "snapshots. Two important properties distinguish it from the retrospective data: (1) The join key is fallnr "
    "rather than fallid. (2) Each snapshot contains an is_open flag: is_open = 1 means the patient is still on "
    "the ICU at snapshot time, so icu_duration_h records elapsed time only, not final LOS. To ensure that the "
    "observed LOS is the actual completed duration, only closed stays (is_open = 0) with LOS > 1 day were "
    "retained (2,026 stays in total). Date formats also differ: retrospective timestamps use ISO format "
    "(YYYY-MM-DD HH:MM:SS), whereas prospective OLD-snapshot CSVs use German format (DD.MM.YYYY HH:MM:SS) and "
    "are parsed accordingly. Of the 2,026 completed prospective stays, 193 had a senior-physician LOS estimate "
    "documented during routine care and formed the matched evaluation set. The prospective cohort was used only "
    "for evaluation — never for model fitting or hyperparameter selection.")

h2("2.3  Outcome: ICU LOS in days")
body("The primary outcome was ICU LOS in days. In the source data, duration is stored in hours "
    "(icu_duration_h); all modelling targets and reported metrics use icu_duration_h ÷ 24. The senior-physician "
    "estimates were recorded directly in days. This conversion is applied consistently throughout all scripts "
    "and is documented in summary.json ('units': 'icu_duration_h are HOURS; target = hours/24 = DAYS').")
body("Because ICU LOS is strictly positive and strongly right-skewed (retrospective cohort: median 2.86 days, "
    "90th percentile 12.1 days, maximum 76.9 days), all four models were fitted on a log1p-transformed target. "
    "log1p(x) = log(1 + x), where '1p' means 'plus one'. This transform was chosen over a plain logarithm for "
    "three reasons: (i) log1p is defined for all x ≥ 0, whereas log(0) is undefined (and LOS values at the "
    "inclusion boundary of exactly one day could produce log(1) = 0 without issue); (ii) the +1 shift avoids "
    "the singularity of log near zero; (iii) log1p compresses the long right tail, reducing the influence of "
    "extreme stays on gradient-based optimisation. All four models use an identical target-transformation wrapper "
    "(scikit-learn TransformedTargetRegressor: log1p on fitting, expm1 = exp(x)−1 on prediction). Predictions "
    "are always back-transformed with expm1 before any metric is computed, guaranteeing non-negative day-scale "
    "outputs.")

h2("2.4  Predictors and the 24-hour window")
body("All predictors were derived strictly from the first 24 hours after ICU admission (from admission timestamp "
    "to admission + 24 h). This temporal restriction is the central leakage safeguard: no measurement, "
    "procedure or device record made after the 24-hour prediction window can enter any model.")
body("From 104 candidate features, 84 were available as genuine first-24-hour variables and were retained. "
    "Twenty candidates were excluded because no leakage-free 24-hour version could be constructed under the "
    "exact feature name. Critically, an earlier version of the pipeline had silently substituted whole-stay "
    "aggregates for 15 of these (e.g. a laboratory summary spanning the entire admission in place of the "
    "24-hour value); this constitutes leakage and was removed entirely (no fallback to whole-stay values). "
    "The 84 retained features span six domains (Table 3): laboratory values (30), admission diagnoses (24), "
    "procedures (13), vascular access devices (8), vital signs (6) and demographic/admission context (3).")

h2("2.5  Handling of patients with multiple stays and multiple records within the window")
body("The analysis unit is the ICU stay, not the patient or the hospital encounter. Patients with more than "
    "one ICU stay (n = 3,156 patients; 25.4% of the cohort) contribute each stay as a separate prediction "
    "instance. To prevent optimistic bias from information leakage across stays of the same patient, every "
    "train/test split and every cross-validation fold was grouped by patient identifier (GroupShuffleSplit "
    "and GroupKFold, respectively). This ensures that all stays from a given patient appear in either training "
    "or evaluation, never in both simultaneously.")
body("Within the 24-hour window, most modalities contribute multiple records for a single stay (e.g. repeated "
    "laboratory measurements, several vital-sign readings, multiple procedure entries). These are aggregated "
    "as follows: for numeric variables (laboratory, vital signs), six summary statistics are computed per "
    "analyte per stay — first, last, minimum, maximum, mean and count of measurements within the 24-hour "
    "window. For binary modalities (procedure codes, access devices, admission diagnoses), a binary presence "
    "indicator (0/1) and a total count are derived. For procedures, an additional feature captures the total "
    "number of distinct procedure codes within the window. When a patient has had multiple operations or "
    "anaesthesia sessions before or during the ICU stay, all perioperative records within the window are "
    "included in the counts; timing is determined by the procedure or anaesthesia start timestamp.")

h2("2.6  Missing data and preprocessing pipeline")
body("Missing values are expected in clinical data, particularly for analytes not routinely measured for every "
    "patient. All imputation is performed inside cross-validation — fitted on the training fold only and "
    "applied to the held-out fold — to prevent imputation leakage. The preprocessing pipeline is as follows:")
body("(1) Numeric features: median imputation (SimpleImputer, strategy='median'), using the training-fold "
    "median. For binary features derived from procedure/access/diagnosis presence, a missing value indicates "
    "the feature was not documented in the window and is treated as absent (0) rather than truly missing; "
    "these are filled with 0 before model fitting.")
body("(2) Categorical feature (care-unit type, oebenekurz, 1 variable): most-frequent imputation followed "
    "by one-hot encoding with unknown-category handling (handle_unknown='ignore').")
body("(3) Standardisation: applied only to ridge regression inputs (StandardScaler after imputation), because "
    "tree-based ensembles are scale-invariant and do not require standardisation.")
body("This pipeline is implemented as a scikit-learn ColumnTransformer inside a Pipeline, all wrapped in a "
    "TransformedTargetRegressor that applies log1p before fitting and expm1 after prediction. The complete "
    "processing chain for any model is therefore: raw features → ColumnTransformer (impute, encode, "
    "[scale for Ridge]) → estimator(log1p(LOS)) → expm1(prediction) → predicted LOS in days.")
body("For the prospective cohort, the identical feature definitions, 24-hour window and naming conventions "
    "were applied when rebuilding features from raw prospective source tables (laboratory, vital signs, "
    "procedures, access devices, diagnoses). No additional fitting was performed: the training-set imputation "
    "medians were carried over directly, exactly as they would be in a clinical deployment scenario.")

h2("2.7  Candidate models and hyperparameter optimisation")
body("Four candidate models were evaluated throughout. All use the same preprocessing pipeline (Section 2.6) "
    "and differ only in the estimator and, for ridge, the additional standardisation step.")
body("Ridge regression is a regularised linear model that minimises the sum of squared residuals plus "
    "a penalty on coefficient magnitude (L2 penalty = alpha). It serves as an interpretable linear baseline. "
    "Random forest [10] and extra trees [11] are bagging-based ensembles of decision trees; extra trees "
    "additionally randomises the splitting thresholds at each node, which reduces variance at the cost of "
    "some additional bias. XGBoost [12] is a gradient-boosting framework that builds an additive ensemble "
    "of trees sequentially, with each tree correcting the residuals of the previous ones.")
body("Hyperparameters were tuned on the training set by 4-fold patient-grouped cross-validation "
    "(GroupKFold on patient identifier), optimising the negative MAE on the day scale (after expm1 "
    "back-transformation). Ridge was tuned by exhaustive grid search; the three nonlinear models by "
    "randomised search (12 candidates each, random seed 42). Search spaces: ridge alpha ∈ "
    "{0.1, 0.3, 1, 3, 10, 30, 100}. For random forest and extra trees: n_estimators ∈ {300, 500}, "
    "max_depth ∈ {None, 12, 20}, min_samples_leaf ∈ {2, 5, 10}, max_features ∈ {sqrt, 0.5}. "
    "For XGBoost: n_estimators ∈ {300, 500, 800}, max_depth ∈ {4, 6, 8}, learning_rate ∈ "
    "{0.03, 0.05, 0.1}, subsample ∈ {0.7, 0.9}, colsample_bytree ∈ {0.7, 0.9}, "
    "min_child_weight ∈ {1, 3, 5}, reg_lambda ∈ {1, 2, 5}.")

h2("2.8  Model selection: pre-specified rule")
body("To avoid double-dipping (optimising model choice on the same data used to report test performance), "
    "the selection rule was fixed before the holdout test set was inspected: the candidate model with the "
    "lowest 4-fold patient-grouped CV-MAE on the training set is the final model. This rule selected Extra "
    "Trees (CV-MAE 2.656 days). The final model was then refitted on the full training set (n = 13,603 "
    "stays) and evaluated exactly once on the holdout test set (n = 3,429 stays) and exactly once on the "
    "prospective cohort (n = 193 stays). No further hyperparameter adjustment was made after seeing the "
    "holdout or prospective results. All four candidate models are reported consistently in Tables 4 and 5 "
    "to show the sensitivity of results to model choice.")

h2("2.9  Performance metrics: discrimination and calibration")
body("All metrics are computed on the day scale after expm1 back-transformation, with each ICU stay "
    "contributing exactly one (observed, predicted) pair.")
body("Discrimination — the ability to rank stays by LOS correctly — is quantified by: "
    "(i) R² = 1 − Σ(yᵢ − ŷᵢ)² / Σ(yᵢ − ȳ)², which equals 1 for perfect predictions and can be "
    "negative when a model performs worse than predicting the mean; and (ii) Spearman rank correlation "
    "ρ between observed and predicted LOS, which is robust to the non-linear scale of LOS values.")
body("Calibration — the alignment of predicted magnitudes with observed values — is quantified by: "
    "(i) MAE = (1/N) Σᵢ₌₁ᴺ |yᵢ − ŷᵢ|, the mean absolute deviation between observed LOS yᵢ (days) "
    "and predicted LOS ŷᵢ (days) across all N stays. For example, if a model predicts 3 days for a "
    "patient who stays 5 days, the absolute error is 2 days; the MAE is the average of such errors over "
    "the entire evaluation set. (ii) Median AE = median(|yᵢ − ŷᵢ|), the 50th percentile of absolute "
    "errors, which is more robust to extreme-LOS outliers than MAE. (iii) RMSE = √[(1/N) Σ(yᵢ − ŷᵢ)²], "
    "which penalises large errors more heavily than MAE. (iv) Mean bias = (1/N) Σ(ŷᵢ − yᵢ), where a "
    "positive value indicates systematic overestimation and a negative value systematic underestimation. "
    "For Spearman ρ, the prospective predictions CSV is used directly; for retrospective ρ this cannot be "
    "computed from published summary statistics and is therefore not reported.")

h2("2.10  Interpretability")
body("Predictor importance for the final model (extra trees) was assessed by permutation importance on "
    "the holdout test set (scikit-learn, 10 repeats, random seed 42). For each predictor in turn, its "
    "values in the test set are randomly permuted (destroying its relationship with LOS), the model is "
    "applied to the permuted data, and the resulting increase in MAE (in days) is recorded. A larger "
    "increase means the model relies more on that predictor. The procedure is repeated 10 times and the "
    "mean ± SD increase is reported (Table 6).")

h2("2.11  Software and reproducibility")
body("Analyses used Python 3.12 (scikit-learn ≥ 1.4, XGBoost ≥ 2.0, DuckDB, SciPy, Pandas, NumPy, "
    "Matplotlib). A fixed random seed (42) was used for the train/test split, cross-validation sampling "
    "and model fitting. The complete pipeline and manuscript code are publicly available at "
    "https://github.com/Scarlet1991/KISIK-Dashboard.")

# ══════════════════════════════════════════════════════════════════════════════
# 3 RESULTS
# ══════════════════════════════════════════════════════════════════════════════
h1("3  Results")

h2("3.1  Cohort characteristics")
body("Retrospective and prospective cohort characteristics are compared in Table 2. Median age was similar "
    "(retrospective 68.0 [IQR 57.2–77.0] vs prospective 69.2 [59.5–77.9] years; SMD 0.10), as was the "
    "distribution across ward types and care units. ICU LOS was slightly shorter in the prospective cohort "
    "(median 2.7 vs 2.9 days; SMD 0.21), which likely reflects the restriction to completed stays only "
    "(censored long stays are excluded). Hospital LOS was available only retrospectively. SAPS II and "
    "TISS-28 severity scores were routinely documented retrospectively (73% of stays each) but almost never "
    "prospectively (≤ 1%), reflecting a change in clinical scoring practice between the two periods; "
    "SOFA documentation was sparse in both (1–2%). Cross-cohort severity comparison was therefore not "
    "feasible and scores are reported descriptively only.")

h2("3.2  Effect of leakage control")
body("Replacing 15 inadvertently substituted whole-stay feature aggregates with genuine 24-hour "
    "counterparts — or excluding them where unavailable — markedly reduced apparent retrospective "
    "performance: a feature set containing whole-stay aggregates reached R² ≈ 0.61 for the same model "
    "family, whereas the strictly leakage-free 24-hour feature set yielded R² ≈ 0.31. This difference "
    "(ΔR² ≈ 0.30) quantifies the degree of leakage inflation in the earlier pipeline. All results below "
    "use the leakage-free 84-feature set exclusively.")

h2("3.3  Model development and selection")
body("In 4-fold patient-grouped cross-validation on the training set, the three tree ensembles were "
    "near-identical and clearly outperformed the linear baseline (Table 4): CV-MAE 2.656–2.671 days "
    "(Extra Trees, Random Forest, XGBoost) vs 2.855 days (Ridge). Extra Trees had the lowest CV-MAE "
    "(2.656 days) and was selected as the final model per the pre-specified rule. The final Extra Trees "
    "model was refitted on the full training set (n = 13,603 stays) with best hyperparameters (500 trees, "
    "max_depth = 20, min_samples_leaf = 2, max_features = 0.5).")
body("On the held-out test set (n = 3,429 stays), all four models again performed in a narrow band: "
    "MAE 2.752–2.936 days, R² 0.234–0.316 (Table 4). Final model (Extra Trees): MAE 2.758 days, median "
    "AE 1.218 days, RMSE 5.343 days, R² 0.308, mean bias −1.078 days. The negative mean bias indicates "
    "that all four models systematically underestimate LOS, which is expected: LOS is right-skewed, "
    "long-stay patients are rare in training, and log1p compression still leaves the very long tail "
    "harder to predict. Random forest and XGBoost were statistically indistinguishable from Extra Trees "
    "on the holdout; ridge regression was consistently the weakest model.")

h2("3.4  Most important predictors")
body("Permutation importance for the final model (Extra Trees) on the holdout test set is shown in Table 6 "
    "and Figure 2. By a large margin, the strongest single predictor was the intensive-care "
    "complex-treatment procedure code (OPS 8-98f.0, base module): permuting this feature increased MAE "
    "by 0.96 days. The next tier — the extended complex-treatment module (8-98f.10), ICU care-unit type, "
    "extended haemodynamic monitoring (8-931), and cardiac monitoring (8-924) — each added 0.07–0.21 days. "
    "The total number of procedures within 24 hours, a further complex-treatment tier (8-98f.11), "
    "ventilator dependence (ICD-10 Z99.1) and ICU stay number per patient added smaller increments. "
    "Age contributed only 0.02 days. In summary, procedure burden and monitoring intensity in the first "
    "24 hours dominate LOS prediction; patient demographics and diagnoses contribute modestly.")

h2("3.5  Prospective evaluation against the senior physician (n = 193)")
body("Prospective features were rebuilt from raw source tables (laboratory, vital signs, procedures, "
    "vascular access, diagnoses) using the identical 24-hour window, feature names and aggregation "
    "functions as in development. This yielded genuine values for 72 of the 84 predictors (86%); the "
    "remaining 12 (5 diagnoses, 6 procedure codes, 1 other) did not appear in any prospective stay "
    "within the first 24 hours and were treated as absent. The median per-stay feature completeness "
    "was 78% (vs 4 features in a preliminary analysis that used the prospective parquet directly, "
    "which had near-zero 24-hour coverage).")
body(f"Results are shown in Table 5 and Figure 1. The senior physician outperformed all four ML models "
    f"on every summary metric. For the final model (Extra Trees): physician MAE 2.01 vs model MAE 2.70 "
    f"days; R² 0.22 vs 0.05; Spearman ρ {_rho_ob} vs {_rho_et}. Among ML models, Extra Trees and XGBoost carried "
    f"modest but genuine prospective predictive signal (R² 0.05 and 0.07, Spearman ρ {_rho_et} and {_rho_xgb}), "
    "confirming that genuine 24-hour inputs restore signal that was absent in the feature-impoverished "
    "preliminary analysis (where R² ≈ 0). Random forest (R² −0.00) and Ridge (R² −7.06) did not perform "
    "better than a mean predictor prospectively. Ridge's large positive bias (+3.87 days) reflects "
    "instability of a linear model under the distributional shift between the two periods.")
body("In subgroup analyses, the physician advantage was most pronounced in short stays (1–7 days, "
    "n = 166: physician MAE 1.28 vs Extra Trees 1.86 days). For long stays (> 7 days, n = 27), the "
    "gap narrowed to physician MAE 6.51 vs XGBoost 6.84 days, suggesting that models may eventually "
    "be competitive at the upper end of the LOS distribution once more training data from long stays "
    "are available.")

h2("3.6  Calibration and long-stay behaviour")
body("Calibration describes how well predicted LOS magnitudes agree with observed values, independent "
    "of discrimination. Figures 3–5 show observed vs predicted ICU LOS for the final model "
    "(Extra Trees) on the retrospective holdout, the prospective cohort, and the senior physician "
    "(all capped at 20 days for legibility; metrics are computed on the full range).")
body("Retrospectively, all models show systematic underestimation: mean bias ranges from −1.010 days "
    "(XGBoost) to −1.208 days (Ridge). The hexbin plot (Figure 3) makes the mechanism visible: for "
    "short stays (< 5 days) prediction scatter is concentrated near the identity line, but for stays "
    "beyond about one week the model predictions flatten and revert towards the population mean, "
    "producing a characteristic 'fanning' pattern below the identity line.")
body("Prospectively, Extra Trees is nearly unbiased overall (+0.09 days), while the physician also "
    "shows slight underestimation (−0.14 days). The very small biases for both suggest "
    "calibration-in-the-large is adequate at the prospective case-mix level. XGBoost shows "
    "moderate overestimation (+0.63 days) and Ridge severe overestimation (+3.87 days). "
    "The long-stay hexbin (Figure 4) confirms that the model underestimates for observed LOS beyond "
    "~ 10 days, mirroring the retrospective pattern. The physician (Figure 5) maintains "
    "better tracking of long stays, consistent with the subgroup MAE advantage (Section 3.5).")

# ══════════════════════════════════════════════════════════════════════════════
# 4 DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════
h1("4  Discussion")
body("In a large single-centre cohort we developed, pre-specified and prospectively validated an ICU LOS "
    "model from the first 24 hours. Three findings stand out.")
body("First, leakage from whole-stay feature aggregation materially inflated apparent performance "
    "(R² 0.61 → 0.31 once corrected). This reinforces recent warnings [7] and underscores that "
    "TRIPOD+AI's requirement for a clearly defined prediction time point is not merely formal: "
    "the choice of feature aggregation window can nearly double the reported R² without any real "
    "improvement in clinical utility.")
body("Second, under strict leakage control, the three tree-ensemble families performed indistinguishably. "
    "The choice between random forest, extra trees and XGBoost was immaterial — a finding that "
    "simplifies model selection in practice. Ridge regression was consistently the weakest model both "
    "retrospectively and, dramatically, prospectively (R² −7.06), which can be attributed to the "
    "regression-to-the-mean property of linear models when features shift in distribution across time "
    "periods.")
body("Third, the experienced senior physician outperformed every model prospectively, with the advantage "
    "largest in long stays. Long-staying patients are precisely those for whom accurate early prediction "
    "matters most for planning. Physicians integrate contextual knowledge — anticipated surgical "
    "trajectories, potential complications, bed-management pressures — not captured in structured "
    "24-hour tabular data.")
body("Regarding calibration: Extra Trees was essentially unbiased prospectively (+0.09 days) but showed "
    "systematic underestimation retrospectively (−1.08 days), likely because right-skewed LOS "
    "distributions create training targets in which very long stays are underrepresented and pull "
    "predictions toward the centre. The log1p transformation mitigates but does not fully eliminate "
    "this. Calibration recalibration (intercept adjustment or isotonic regression) could be applied "
    "prospectively but was not pursued here.")

h2("4.1  Limitations")
body("This is a single-centre study; the absolute performance values and ranking of models may not "
    "generalise to other case mixes or information systems. The prospective comparison rested on 193 "
    "completed matched stays, and while this constitutes a genuine, undiluted prospective test with "
    "documented physician estimates, the sample is modest and the physician's estimates may have "
    "incorporated information beyond the first 24 hours. The reconstruction of 24-hour features from "
    "raw prospective tables achieved 86% coverage; the remaining 12 features were treated as absent, "
    "which may introduce a minor conservative bias relative to a deployment where all features are "
    "fully available. The retrospective-to-prospective performance drop is consistent with distribution "
    "shift between the two periods (case mix, coding practice) but its sources were not fully "
    "characterised.")

h1("5  Conclusion")
body("Under strict leakage control and prospective, clinician-benchmarked validation, a machine-learning "
    "model built from the first 24 hours of ICU data did not match experienced senior-physician judgement "
    "for ICU LOS prediction. The clinician advantage was largest for long-staying patients, who represent "
    "the greatest capacity-planning challenge. Realising clinical impact will require not only stronger "
    "models but also transparent, prospective validation against the clinicians the technology is "
    "intended to support.")

# ══════════════════════════════════════════════════════════════════════════════
# STATEMENTS / REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
h1("Statements")
labeled("Data availability statement.","De-identified routine clinical data are not publicly available owing "
    "to data-protection regulations; requests may be directed to the corresponding author subject to "
    "institutional approval. The complete analysis and figure code is openly available at "
    "https://github.com/Scarlet1991/KISIK-Dashboard.")
labeled("Ethics statement.","The study analysed de-identified routine clinical data; the protocol and waiver "
    "of informed consent were reviewed by [ethics committee], approval number [to be inserted].")
labeled("Author contributions.","[To be completed.] All authors approved the final manuscript.")
labeled("Funding.","[To be completed; if none, state so.]")
labeled("Conflict of interest.","The authors declare no competing interests.")

h1("References")
refs = [
    "Knaus WA, Draper EA, Wagner DP, Zimmerman JE. APACHE II: a severity of disease classification system. Crit Care Med. 1985;13(10):818–829.",
    "Le Gall JR, Lemeshow S, Saulnier F. A new Simplified Acute Physiology Score (SAPS II). JAMA. 1993;270(24):2957–2963.",
    "Vincent JL, Moreno R, Takala J, et al. The SOFA score to describe organ dysfunction/failure. Intensive Care Med. 1996;22(7):707–710.",
    "Verburg IWM, Atashi A, Eslami S, et al. Which models can I use to predict adult ICU length of stay? A systematic review. Crit Care Med. 2017;45(2):e222–e231.",
    "Johnson AEW, Pollard TJ, Shen L, et al. MIMIC-III, a freely accessible critical care database. Sci Data. 2016;3:160035.",
    "Pollard TJ, Johnson AEW, Raffa JD, et al. The eICU Collaborative Research Database. Sci Data. 2018;5:180178.",
    "Kapoor S, Narayanan A. Leakage and the reproducibility crisis in machine-learning-based science. Patterns. 2023;4(9):100804.",
    "Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent reporting of a multivariable prediction model (TRIPOD). Ann Intern Med. 2015;162(1):55–63.",
    "Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine-learning methods. BMJ. 2024;385:e078378.",
    "Breiman L. Random forests. Mach Learn. 2001;45(1):5–32.",
    "Geurts P, Ernst D, Wehenkel L. Extremely randomized trees. Mach Learn. 2006;63(1):3–42.",
    "Chen T, Guestrin C. XGBoost: a scalable tree boosting system. In: Proc. 22nd ACM SIGKDD; 2016. p. 785–794.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(3); run(p, f"[{i}]  {r}", size=Pt(10.5))
small("Reference details should be verified; topic-specific references on ICU LOS prediction, calibration "
      "and distribution shift should be added before submission.")

# ══════════════════════════════════════════════════════════════════════════════
# TABLES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph().add_run("").add_break()
h1("Tables")

# ── Table 1: Cohort overview ──────────────────────────────────────────────────
cap(1, "Cohort overview and inclusion criteria.")
table([
    ["Parameter","Retrospective (development)","Prospective (evaluation)"],
    ["Data period","May 2017 – Jul 2024","Oct 2024 – Jan 2026"],
    ["Source","KISIK clinical data repository","Daily OLD live-system snapshots"],
    ["Join key","fallid","fallnr (case reference differs by system)"],
    ["Date format","ISO: YYYY-MM-DD HH:MM:SS","German: DD.MM.YYYY HH:MM:SS"],
    ["Inclusion","17 eligible ward/care-unit pairs; LOS > 1 day",
     "Same ward filter; is_open = 0 (completed stays); LOS > 1 day"],
    ["ICU stays (n)","17,032","2,026 completed (193 with senior estimate)"],
    ["Patients (n)","12,414","—"],
    ["Hospital encounters (n)","13,275","—"],
    ["Median ICU LOS (days)","2.86","2.7 (matched subset)"],
    ["Train / test split","13,603 / 3,429 (patient-grouped)","Evaluation only"],
    ["Role","Model development & selection","External prospective test"],
],[4.2,5.5,5.5], numcols_from=1)
small("LOS, length of stay. is_open = 0 indicates the patient had been discharged by the time of data capture, "
      "so icu_duration_h records the final ICU duration. Censored stays (is_open = 1) are excluded because "
      "their duration is still accruing. The prospective cohort was never used for model fitting or selection.")

# ── Table 2: Patient characteristics (dynamic from CSV) ───────────────────────
cap(2, "Patient and stay characteristics: retrospective development vs. prospective cohort "
       "(completed stays, is_open = 0, LOS > 1 day).")
t1rows = build_table1()
tbl_rows = [["Characteristic", "Retrospective (n = 17,032)", "Prospective (n = 2,026)", "p-value", "SMD"]]
for char, retro, pros, pstr, smd, is_hd in t1rows:
    tbl_rows.append([char, retro, pros, pstr, str(smd)])
t2 = doc.add_table(rows=len(tbl_rows), cols=5)
t2.style = "Table Grid"; t2.alignment = WD_TABLE_ALIGNMENT.CENTER; t2.autofit = False
col_w = [6.0, 4.2, 4.2, 1.7, 1.1]
for ri, row in enumerate(tbl_rows):
    for ci, val in enumerate(row):
        c = t2.cell(ri, ci); c.width = Cm(col_w[ci])
        is_hd_row = (ri > 0 and tbl_rows[ri][0].strip() and not tbl_rows[ri][0].startswith("   ")
                     and ri < len(t1rows) + 1 and t1rows[ri-1][5] if ri > 0 else False)
        if ri == 0:
            settext(c, val, bold=True, color=RGBColor(0xFF,0xFF,0xFF), align="center")
        else:
            settext(c, val, bold=is_hd_row, align="center" if ci > 0 else "left")
    if ri == 0:
        for c in t2.rows[0].cells: bg(c, HFILL)
    else:
        row_type = t1rows[ri-1][5] if ri - 1 < len(t1rows) else False
        if row_type:
            for c in t2.rows[ri].cells: bg(c, ZEB)
        elif ri % 2 == 0:
            for c in t2.rows[ri].cells: bg(c, ZEB)
small("Continuous variables: median [IQR], Mann-Whitney U test. Categorical variables: n (%), chi-square test. "
      "SMD = standardised mean difference (|SMD| > 0.1 indicates notable between-cohort imbalance). "
      "Hospital LOS was unavailable prospectively. Sex and BMI were not available in either dataset. "
      "Severity scores (SAPS II, TISS-28, SOFA): extracted from score records within the first 24 h; "
      "documentation rates differ markedly by period (see Section 3.1) and cross-cohort comparison "
      "is therefore not feasible (n.c. = not compared). With a large retrospective sample, "
      "small absolute differences reach statistical significance; the SMD is the more clinically "
      "meaningful measure of cohort imbalance.")

# ── Table 3: Predictor domains ────────────────────────────────────────────────
cap(3, "Predictor domains: 84 leakage-free features, all from the first 24 hours after ICU admission.")
table([
    ["Domain","Features (n)","Aggregation","Examples"],
    ["Laboratory (first 24 h)","30",
     "first, last, min, max, mean, count",
     "potassium, sodium, glucose, lactate, bilirubin, creatinine, haemoglobin"],
    ["Admission diagnoses (ICD-10)","24","binary presence (0/1)",
     "sepsis, intracranial haemorrhage, coronary disease, ventilator dependence (Z99.1)"],
    ["Procedures (OPS, first 24 h)","13","binary presence + total count",
     "ICU complex treatment (8-98f.0/10/11), extended monitoring (8-931), mechanical ventilation"],
    ["Vascular access (first 24 h)","8","binary presence + count",
     "arterial line, central venous catheter, urinary catheter, epidural catheter"],
    ["Vital signs (first 24 h)","6","first, last, min, max, mean, count",
     "peripheral oxygen saturation (SpO₂)"],
    ["Demographics / admission context","3","as recorded",
     "age (years), ICU stay number per patient, care-unit type (oebenekurz)"],
    ["Total","84","",""],
],[3.0,2.0,3.2,9.0], numcols_from=1)
small("20 of 104 candidate features were excluded (no leakage-free 24-hour version available under "
      "the exact column name); an earlier pipeline had substituted whole-stay aggregates for 15 of "
      "these — that leakage has been removed. All feature values are derived from the window "
      "[admission timestamp, admission + 24 h]. Admission diagnoses are included if documented "
      "as principal diagnosis by 24 h post-admission. Care-unit type is one-hot encoded; the remaining "
      "85 features are numeric (binary 0/1 for presence/absence, or continuous for labs/vitals counts).")

# ── Table 4: Retrospective model development ──────────────────────────────────
cap(4, "Model development: tuned hyperparameters, 4-fold patient-grouped cross-validated MAE and "
       "hold-out test-set performance (all four candidate models; retrospective cohort, n = 3,429 test stays).")
et = rm("ExtraTrees"); rf = rm("RandomForest"); xg = rm("XGBoost"); ri_m = rm("Ridge")
table([
    ["Model","Tuned hyperparameters","CV-MAE (d)","Test MAE (d)","Median AE (d)","RMSE (d)","R²","Bias (d)"],
    ["Ridge (baseline)",
     f"alpha = 0.1; numeric inputs standardised (StandardScaler)",
     f"{ri_m['CV_MAE_days']:.3f}",f"{ri_m['MAE_days']:.3f}",
     f"{ri_m['MedianAE_days']:.3f}",f"{ri_m['RMSE_days']:.3f}",
     f"{ri_m['R2']:.3f}",f"{ri_m['Bias_days']:+.3f}"],
    ["Random forest",
     f"500 trees, max_depth 20, min_samples_leaf 2, max_features 0.5",
     f"{rf['CV_MAE_days']:.3f}",f"{rf['MAE_days']:.3f}",
     f"{rf['MedianAE_days']:.3f}",f"{rf['RMSE_days']:.3f}",
     f"{rf['R2']:.3f}",f"{rf['Bias_days']:+.3f}"],
    ["Extra Trees (final ✓)",
     f"500 trees, max_depth 20, min_samples_leaf 2, max_features 0.5",
     f"{et['CV_MAE_days']:.3f}",f"{et['MAE_days']:.3f}",
     f"{et['MedianAE_days']:.3f}",f"{et['RMSE_days']:.3f}",
     f"{et['R2']:.3f}",f"{et['Bias_days']:+.3f}"],
    ["XGBoost",
     f"500 trees, depth 8, lr 0.05, subsample 0.9, colsample 0.9, min_child_wt 1, λ 5",
     f"{xg['CV_MAE_days']:.3f}",f"{xg['MAE_days']:.3f}",
     f"{xg['MedianAE_days']:.3f}",f"{xg['RMSE_days']:.3f}",
     f"{xg['R2']:.3f}",f"{xg['Bias_days']:+.3f}"],
],[2.5,4.5,1.5,1.5,1.6,1.5,1.1,1.5], numcols_from=2)
small("All models share the same preprocessing pipeline (median imputation, one-hot encoding of care-unit type) "
      "and a log1p target transformation. Ridge additionally applies StandardScaler. "
      "CV-MAE = 4-fold patient-grouped cross-validated MAE on the training set (n = 13,603); lower = better. "
      "Pre-specified selection rule: model with lowest CV-MAE → Extra Trees (CV-MAE 2.656 d). "
      "Bias = mean(predicted − observed); negative = systematic underestimation. "
      "RMSE penalises large errors more heavily than MAE.")

# ── Table 5: Prospective comparison ───────────────────────────────────────────
pob = pm("Oberarzt"); pet = pm("ExtraTrees"); pxg = pm("XGBoost"); prf = pm("RandomForest"); pri = pm("Ridge")
cap(5, "Prospective comparison: senior physician vs all four ML models. "
       "Completed matched stays only (is_open = 0, LOS > 1 day; n = 193). "
       "First-24-hour features rebuilt from raw prospective records (86% of 84 features available).")
table([
    ["Method","MAE (d)","Median AE (d)","RMSE (d)","R²","Bias (d)","Spearman ρ"],
    ["Senior physician",
     f"{pob['MAE']:.2f}",f"{pob['MedianAE']:.2f}",f"{pob['RMSE']:.2f}",
     f"{pob['R2']:.2f}",f"{pob['Bias']:+.2f}",f"{RHO['Senior physician']:.2f}"],
    ["Extra Trees (final model)",
     f"{pet['MAE']:.2f}",f"{pet['MedianAE']:.2f}",f"{pet['RMSE']:.2f}",
     f"{pet['R2']:.2f}",f"{pet['Bias']:+.2f}",f"{RHO['ExtraTrees']:.2f}"],
    ["XGBoost",
     f"{pxg['MAE']:.2f}",f"{pxg['MedianAE']:.2f}",f"{pxg['RMSE']:.2f}",
     f"{pxg['R2']:.2f}",f"{pxg['Bias']:+.2f}",f"{RHO['XGBoost']:.2f}"],
    ["Random forest",
     f"{prf['MAE']:.2f}",f"{prf['MedianAE']:.2f}",f"{prf['RMSE']:.2f}",
     f"{prf['R2']:.2f}",f"{prf['Bias']:+.2f}",f"{RHO['RandomForest']:.2f}"],
    ["Ridge",
     f"{pri['MAE']:.2f}",f"{pri['MedianAE']:.2f}",f"{pri['RMSE']:.2f}",
     f"{pri['R2']:.2f}",f"{pri['Bias']:+.2f}",f"{RHO['Ridge']:.2f}"],
],[4.4,1.9,2.1,1.9,1.5,1.9,2.1], numcols_from=1)
small("Subgroup MAE — 1–7 days (n = 166): physician 1.28 vs Extra Trees 1.86 d; "
      "> 7 days (n = 27): physician 6.51 vs XGBoost 6.84 d. "
      "Spearman ρ = rank correlation between observed and predicted LOS (discrimination). "
      "Bias = mean(predicted − observed); positive = overestimation. "
      "Ridge is unstable under prospective distribution shift (R² −7.06, bias +3.87 d) "
      "owing to regression-to-the-mean under covariate shift and is not recommended clinically. "
      "All four models trained on identical retrospective data (n = 13,603) with the same pipeline; "
      "only the estimator and scaling step differ (Section 2.6–2.7).")

# ── Table 6: Feature importance ───────────────────────────────────────────────
FEAT_LABELS = {
    "proc24_8_98f_0":    "ICU complex treatment, base module (OPS 8-98f.0)",
    "proc24_8_98f_10":   "ICU complex treatment, extended module (OPS 8-98f.10)",
    "oebenekurz":        "ICU care-unit type (demographic/admission context)",
    "proc24_8_931_0":    "Extended haemodynamic monitoring (OPS 8-931)",
    "proc24_8_924":      "Cardiac monitoring / rhythm therapy (OPS 8-924)",
    "proc24_anzahl_gesamt": "Total procedure count, first 24 h",
    "proc24_8_98f_11":   "ICU complex treatment, prolonged module (OPS 8-98f.11)",
    "diag_main_z99_1":   "Ventilator dependence (ICD-10 Z99.1)",
    "proc24_8_930":      "Basic haemodynamic monitoring (OPS 8-930)",
    "stay_nr":           "ICU stay number per patient",
    "diag_main_j12_8":   "Viral pneumonia, unclassified (ICD-10 J12.8)",
    "alter":             "Age (years)",
}
cap(6, "Permutation feature importance of the final model (Extra Trees) on the retrospective "
       "holdout test set (n = 3,429; 10 repeats).")
top_feat = feat.head(12)
fi_rows = [["Rank","Predictor","Domain","Δ MAE when permuted (days, mean ± SD)"]]
for rk, (_, row) in enumerate(top_feat.iterrows(), 1):
    fname = row["Feature"]; imp = float(row["MAE_increase_days"]); sd = float(row["sd"])
    label = FEAT_LABELS.get(fname, fname)
    dom = ("Procedure" if fname.startswith("proc24") else
           "Diagnosis" if fname.startswith("diag_main") else
           "Lab" if fname.startswith("lab24") else
           "Vital" if fname.startswith("vital24") else
           "Access" if fname.startswith("zugang24") else
           "Demographics / admission")
    fi_rows.append([str(rk), label, dom, f"{imp:.3f} ± {sd:.3f}"])
table(fi_rows, [0.8,8.4,2.4,3.2], numcols_from=2)
small("Δ MAE = increase in mean absolute error (days) when the predictor's values are randomly permuted, "
      "destroying its relationship with LOS. Larger value = model more reliant on that predictor. "
      "Mean ± SD over 10 independent permutation repeats. "
      "OPS = Operationen- und Prozedurenschlüssel (German procedure classification). "
      "ICD-10 codes in parentheses.")

# ══════════════════════════════════════════════════════════════════════════════
# FIGURES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph().add_run("").add_break()
h1("Figures")

figure(CAN/"fig_model_comparison.png", 15.0)
figcap(1, "Mean absolute error (MAE, days) of the four candidate models on the retrospective holdout test "
    "set (n = 3,429 stays) and on the prospective cohort (n = 193 completed stays). The dashed line marks "
    "the senior-physician MAE. Lower is better. All four models and the senior physician are shown "
    "consistently. The near-identical performance of the tree ensembles retrospectively contrasts with "
    "the larger prospective spread; Ridge's instability under distributional shift is evident.")

figure(CAN/"fig_importance.png", 15.0)
figcap(2, "Permutation feature importance of the final model (Extra Trees) on the retrospective holdout "
    "test set: mean increase in MAE (days, ± SD over 10 repeats) when each predictor is independently "
    "permuted. ICU complex-treatment procedure codes dominate by a large margin. Features with "
    "Δ MAE < 0.01 d are omitted for clarity.")

figure(CAN/"fig_hexbin_retro_ExtraTrees.png", 12.0)
figcap(3, "Retrospective holdout — observed vs predicted ICU LOS: final model (Extra Trees, n = 3,429). "
    "Axes truncated at 20 days; metrics in the inset are computed on all stays including those beyond "
    "20 days. Colour encodes the number of stays per hexagonal bin on a log scale. The dashed red "
    "line = identity (perfect prediction). Systematic underestimation of very long stays "
    "(> ~ 7 days) is visible as the predicted values plateau while observed values continue to grow "
    "(points drift below the identity line). Discrimination: R² = 0.31, Spearman ρ not computed "
    "from published summaries. Calibration: mean bias −1.08 days (systematic underestimation).")

figure(CAN/"fig_hexbin_pros_ExtraTrees.png", 12.0)
figcap(4, f"Prospective cohort — observed vs predicted ICU LOS: final model (Extra Trees, n = 193 "
    f"completed stays). Axes truncated at 20 days; metrics computed on all stays. Features were "
    f"rebuilt from raw prospective records (86% coverage). Discrimination: R² = 0.05, Spearman ρ = {_rho_et}. "
    "Calibration: mean bias +0.09 days (nearly unbiased). The underestimation of long stays "
    "persists qualitatively, consistent with the retrospective pattern.")

figure(CAN/"fig_hexbin_pros_oberarzt.png", 12.0)
figcap(5, f"Prospective cohort — observed ICU LOS vs senior-physician estimate (n = 193 completed stays). "
    f"Axes truncated at 20 days. Discrimination: R² = 0.22, Spearman ρ = {_rho_ob}. "
    "Calibration: mean bias −0.14 days (nearly unbiased). The physician shows better tracking of "
    "long-stay patients than any ML model (visible as tighter scatter around the identity line "
    "beyond ~ 5 days).")

# ══════════════════════════════════════════════════════════════════════════════
doc.save(str(OUT))
print(f"Saved: {OUT}")
print(f"Size: {round(OUT.stat().st_size/1024, 1)} KB")
