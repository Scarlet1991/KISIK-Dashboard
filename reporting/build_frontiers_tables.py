# -*- coding: utf-8 -*-
"""
Erzeugt ein publikationsfertiges Word-Dokument mit drei Tabellen
fuer die Frontiers-Einreichung (Research Topic: MedicineAI).
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

OUT = Path(r"D:\Ausgangsdaten\KISIK Projekt\Eigene Auswertung\KISIK_Frontiers_Tabellen.docx")

HEADER_FILL = "1F4E79"   # dunkelblau
HEADER_TXT  = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA_FILL  = "EAF1F8"   # hellblau
BOLD_FILL   = "D9E6F2"   # fuer Hervorhebung bestes Modell

FONT = "Calibri"

# --------------------------------------------------------------------------
doc = Document()

# Grundschrift
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10)

sec = doc.sections[0]
sec.orientation = WD_ORIENT.PORTRAIT
# A4
sec.page_width  = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.0)
sec.top_margin = sec.bottom_margin = Cm(2.0)

CONTENT_W = Cm(21.0 - 4.0)  # 17 cm nutzbar


# --------------------------------------------------------------------------
# Hilfsfunktionen
# --------------------------------------------------------------------------
def set_cell_bg(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, align="left", italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, col_widths_cm, header_rows=1, zebra=True,
                num_align_from_col=None, highlight_row=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    # Spaltenbreiten
    for r in table.rows:
        for idx, cell in enumerate(r.cells):
            cell.width = Cm(col_widths_cm[idx])
    # Header faerben
    for hr in range(header_rows):
        for cell in table.rows[hr].cells:
            set_cell_bg(cell, HEADER_FILL)
    # Zebra
    if zebra:
        for ri in range(header_rows, len(table.rows)):
            if (ri - header_rows) % 2 == 1:
                for cell in table.rows[ri].cells:
                    set_cell_bg(cell, ZEBRA_FILL)
    # Hervorgehobene Zeile
    if highlight_row is not None:
        for cell in table.rows[highlight_row].cells:
            set_cell_bg(cell, BOLD_FILL)


def add_caption(num, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"Table {num}. ")
    r.font.bold = True
    r.font.size = Pt(10.5)
    r.font.name = FONT
    r2 = p.add_run(title)
    r2.font.size = Pt(10.5)
    r2.font.name = FONT


def add_footnote(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.name = FONT
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def fill_table(table, data, header_rows=1, num_align_from_col=None):
    """data: Liste von Zeilen (Liste von Strings)."""
    for ri, row in enumerate(data):
        is_header = ri < header_rows
        for ci, val in enumerate(row):
            cell = table.cell(ri, ci)
            align = "left"
            if num_align_from_col is not None and ci >= num_align_from_col:
                align = "center"
            if is_header:
                set_cell_text(cell, val, bold=True, color=HEADER_TXT,
                              size=9.5, align="center")
            else:
                set_cell_text(cell, val, size=9.5, align=align)


# --------------------------------------------------------------------------
# Titel
# --------------------------------------------------------------------------
h = doc.add_paragraph()
hr = h.add_run("Supplementary Tables")
hr.font.bold = True
hr.font.size = Pt(15)
hr.font.name = FONT
hr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sr = sub.add_run("Prediction of intensive-care length of stay from the first 24 hours: "
                 "machine-learning models versus senior-physician estimates")
sr.font.size = Pt(11)
sr.font.italic = True
sr.font.name = FONT


# ==========================================================================
# TABLE 1A - Daten, Kohorte, Verarbeitung
# ==========================================================================
add_caption("1", "Data sources, cohort definition, and processing.")

t1_data = [
    ["Item", "Description"],
    ["Data source",
     "Single tertiary-care centre clinical data repository (KISIK); de-identified routine-care data."],
    ["Linked data modalities",
     "ICU stay records, ICD-10 diagnoses, laboratory results, vital signs, OPS procedure codes, "
     "and vascular-access devices, linked by case identifier."],
    ["Prediction target (outcome)",
     "ICU length of stay (days), modelled as a continuous variable."],
    ["Feature time window",
     "All predictors restricted to the first 24 h after ICU admission "
     "(admission timestamp to admission + 24 h) to prevent outcome leakage."],
    ["Inclusion criteria",
     "ICU care units only (17 ward / care-unit combinations, predominantly anaesthesiological ICU); "
     "ICU stay > 24 h (> 1 day)."],
    ["Retrospective development cohort",
     "17,032 ICU stays."],
    ["Train / test split",
     "Patient-level split (GroupShuffleSplit, 80 / 20); no patient appears in both sets "
     "-> 13,603 training and 3,429 test stays."],
    ["Candidate predictors",
     "104 selected features; 102 available and used; 2 excluded (not present in source data)."],
    ["Prospective validation cohort",
     "5,655 ICU stays (independent, later time period; “OLD” data extract)."],
    ["Senior-physician estimates",
     "360 prospectively documented length-of-stay estimates."],
    ["Matched evaluation set",
     "359 stays with both a model prediction and a senior-physician estimate "
     "(matched on stay identifier)."],
]
t1 = doc.add_table(rows=len(t1_data), cols=2)
t1.style = "Table Grid"
fill_table(t1, t1_data, header_rows=1)
style_table(t1, [4.3, 12.7], header_rows=1, zebra=True)

add_footnote("ICU, intensive care unit; ICD-10, International Classification of Diseases, 10th revision; "
             "OPS, German procedure classification (Operationen- und Prozedurenschluessel).")


# ==========================================================================
# TABLE 1B - Praediktor-Domaenen
# ==========================================================================
add_caption("1B", "Predictor domains (all derived from the first 24 hours after admission).")

t1b_data = [
    ["Domain", "Features (n)", "Representation", "Examples"],
    ["Demographics & admission context", "6", "Numeric / categorical",
     "Age, ICU stay number, admission hour / weekday / month, care-unit type"],
    ["Laboratory values", "42", "Numeric (first, mean, min, max, last, count)",
     "Haemoglobin, bilirubin, potassium, sodium, lactate, glucose"],
    ["Vital signs", "6", "Numeric (first, mean, min, max, last, count)",
     "Peripheral oxygen saturation (SpO₂)"],
    ["Procedures (OPS)", "13", "Binary presence + count",
     "Mechanical ventilation, dialysis, complex ICU monitoring"],
    ["Vascular access", "13", "Binary presence + count",
     "Arterial line, central venous catheter, urinary catheter"],
    ["Principal diagnoses (ICD-10)", "24", "Binary",
     "Coronary artery disease, sepsis, intracranial haemorrhage"],
    ["Total", "104", "", ""],
]
t1b = doc.add_table(rows=len(t1b_data), cols=4)
t1b.style = "Table Grid"
fill_table(t1b, t1b_data, header_rows=1, num_align_from_col=1)
style_table(t1b, [4.2, 1.6, 4.4, 6.8], header_rows=1, zebra=True,
            highlight_row=len(t1b_data) - 1)
# "Total"-Zeile fett
for cell in t1b.rows[-1].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True


# ==========================================================================
# TABLE 2 - Modelle & retrospektive Performance
# ==========================================================================
doc.add_paragraph().add_run("").add_break()
add_caption("2", "Machine-learning models and retrospective hold-out performance "
                 "(patient-level test set, n = 3,429 stays).")

t2_data = [
    ["Model", "MAE (d)", "Median AE (d)", "RMSE (d)", "R²", "Mean bias (d)"],
    ["XGBoost (gradient boosting)", "2.11", "0.92", "4.22", "0.568", "−0.61"],
    ["Ridge regression",           "2.56", "1.15", "5.10", "0.370", "−0.70"],
    ["Random forest",              "2.55", "1.12", "5.14", "0.360", "−1.14"],
    ["Extra trees",                "2.70", "1.20", "5.38", "0.298", "−1.24"],
]
t2 = doc.add_table(rows=len(t2_data), cols=6)
t2.style = "Table Grid"
fill_table(t2, t2_data, header_rows=1, num_align_from_col=1)
style_table(t2, [5.0, 2.4, 2.4, 2.4, 2.4, 2.4], header_rows=1, zebra=True,
            highlight_row=1)
# bestes Modell (Zeile 1) fett
for cell in t2.rows[1].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

add_footnote("All models predict ICU length of stay on a log1p-transformed target "
             "(back-transformed via expm1). Models ranked by R². Best model highlighted. "
             "MAE, mean absolute error; AE, absolute error; RMSE, root-mean-square error; "
             "R², coefficient of determination; bias = mean(predicted − observed); d, days.")


# ==========================================================================
# TABLE 3A - Prospektiver Vergleich ML vs Oberarzt (gesamt)
# ==========================================================================
doc.add_paragraph().add_run("").add_break()
add_caption("3", "Prospective comparison of machine-learning models versus senior-physician "
                 "estimate (matched cohort, n = 359 stays).")

# Unterzeile A
pa = doc.add_paragraph()
pa.paragraph_format.space_after = Pt(3)
ra = pa.add_run("(A) Overall predictive accuracy")
ra.font.bold = True
ra.font.size = Pt(9.5)
ra.font.name = FONT

t3a_data = [
    ["Method", "MAE (d)", "Median AE (d)", "RMSE (d)", "R²", "Bias (d)", "Wilcoxon p¹"],
    ["Senior physician", "2.60", "0.93", "5.25", "0.251", "−0.63", "reference"],
    ["XGBoost",          "3.65", "1.34", "6.63", "−0.20", "−2.29", "< 0.0001"],
    ["Random forest",    "3.57", "1.63", "6.33", "−0.09", "−1.74", "< 0.0001"],
    ["Extra trees",      "3.58", "2.08", "6.11", "−0.01", "−1.21", "< 0.0001"],
    ["Ridge regression", "4.06", "2.11", "6.79", "−0.25", "−0.56", "< 0.0001"],
]
t3a = doc.add_table(rows=len(t3a_data), cols=7)
t3a.style = "Table Grid"
fill_table(t3a, t3a_data, header_rows=1, num_align_from_col=1)
style_table(t3a, [3.4, 2.0, 2.4, 2.0, 1.7, 2.0, 3.5], header_rows=1, zebra=True,
            highlight_row=1)
# Oberarzt-Zeile (Referenz) fett
for cell in t3a.rows[1].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

# Unterzeile B
pb = doc.add_paragraph()
pb.paragraph_format.space_before = Pt(8)
pb.paragraph_format.space_after = Pt(3)
rb = pb.add_run("(B) Mean absolute error (days) by observed length-of-stay subgroup")
rb.font.bold = True
rb.font.size = Pt(9.5)
rb.font.name = FONT

t3b_data = [
    ["LoS subgroup", "n", "Senior", "XGBoost", "Random forest", "Extra trees", "Ridge"],
    ["1–7 days",  "216", "1.39", "1.61", "1.51", "1.57", "2.23"],
    ["> 7 days",       "70",  "7.74", "12.23", "11.60", "10.85", "11.00"],
    ["> 14 days",      "30",  "11.77", "18.89", "18.17", "17.52", "17.19"],
]
t3b = doc.add_table(rows=len(t3b_data), cols=7)
t3b.style = "Table Grid"
fill_table(t3b, t3b_data, header_rows=1, num_align_from_col=1)
style_table(t3b, [3.0, 1.4, 2.0, 2.2, 3.0, 2.2, 2.2], header_rows=1, zebra=True)

# Unterzeile C
pc = doc.add_paragraph()
pc.paragraph_format.space_before = Pt(8)
pc.paragraph_format.space_after = Pt(3)
rc = pc.add_run("(C) Stay-level head-to-head: share of stays predicted more accurately")
rc.font.bold = True
rc.font.size = Pt(9.5)
rc.font.name = FONT

t3c_data = [
    ["Comparison", "ML more accurate (%)", "Senior more accurate (%)"],
    ["XGBoost vs. senior",       "28.7", "71.3"],
    ["Random forest vs. senior", "28.7", "71.3"],
    ["Extra trees vs. senior",   "29.8", "70.2"],
    ["Ridge vs. senior",         "25.9", "74.1"],
]
t3c = doc.add_table(rows=len(t3c_data), cols=3)
t3c.style = "Table Grid"
fill_table(t3c, t3c_data, header_rows=1, num_align_from_col=1)
style_table(t3c, [5.0, 6.0, 6.0], header_rows=1, zebra=True)

add_footnote("¹ Wilcoxon signed-rank test on paired absolute errors (each model versus the "
             "senior physician); all comparisons p < 0.0001 in favour of the senior physician. "
             "Negative R² indicates worse fit than the cohort mean. Subgroups defined by "
             "observed ICU length of stay. LoS, length of stay; d, days.")


# ==========================================================================
# FIGURES
# ==========================================================================
from docx.enum.text import WD_BREAK

FIG_DIR = Path(r"D:\Ausgangsdaten\KISIK Projekt\Eigene Auswertung")

def add_figure(num, filename, caption, width_cm=16.0):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(10)
    p_img.paragraph_format.space_after = Pt(2)
    run = p_img.add_run()
    run.add_picture(str(FIG_DIR / filename), width=Cm(width_cm))
    # Bildunterschrift
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_cap.paragraph_format.space_after = Pt(10)
    r1 = p_cap.add_run(f"Figure {num}. ")
    r1.font.bold = True
    r1.font.size = Pt(9.5)
    r1.font.name = FONT
    r2 = p_cap.add_run(caption)
    r2.font.size = Pt(9.5)
    r2.font.name = FONT

# Seitenumbruch vor dem Abbildungsteil
pb = doc.add_paragraph()
pb.add_run().add_break(WD_BREAK.PAGE)
ph = doc.add_paragraph()
phr = ph.add_run("Figures")
phr.font.bold = True
phr.font.size = Pt(15)
phr.font.name = FONT
phr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

add_figure(1, "fig1_holdout_mae_rmse.png",
           "Retrospective hold-out performance: mean absolute error (MAE) and root-mean-square "
           "error (RMSE) per model (24-hour feature set; patient-level test set, n = 3,429 stays).")
add_figure(2, "fig2_subgroup_mae.png",
           "Prospective mean absolute error by observed length-of-stay subgroup: senior physician "
           "versus machine-learning models (matched cohort, n = 359 stays).")
add_figure(3, "fig3_bland_altman.png",
           "Bland–Altman analysis of systematic over- and under-estimation of ICU length of stay "
           "(best machine-learning model versus senior physician). Solid line, mean bias; "
           "dashed lines, ±1.96 SD limits of agreement.")
add_figure(4, "fig4_scatter_observed_predicted.png",
           "Observed versus predicted ICU length of stay (machine learning versus senior physician). "
           "Dashed line denotes the line of identity.")
add_figure(5, "fig5_calibration.png",
           "Calibration: mean predicted versus observed length of stay per length-of-stay stratum "
           "(prospective cohort).")
add_figure(6, "fig6_staywise_h2h.png",
           "Stay-level head-to-head: proportion of stays predicted more accurately by each "
           "machine-learning model versus the senior physician.")


# --------------------------------------------------------------------------
doc.save(str(OUT))
print(f"Gespeichert: {OUT}")
print(f"Groesse: {round(OUT.stat().st_size/1024, 1)} KB")
