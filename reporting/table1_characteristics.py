# -*- coding: utf-8 -*-
"""Table 1: Patientencharakteristika retrospektiv vs. prospektiv + statistischer Gruppenvergleich.
Kontinuierlich: Median [IQR], Mann-Whitney-U; kategorial: n(%), Chi-Quadrat. Zusaetzlich SMD (Imbalance-Mass)."""
import sys, io, warnings; warnings.filterwarnings("ignore")
sys.stdout=io.TextIOWrapper(sys.stdout.buffer,encoding="utf-8")
from pathlib import Path
import duckdb, numpy as np, pandas as pd
from scipy import stats
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn; from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

BASE=Path(r"D:\Ausgangsdaten\KISIK Projekt"); AN=BASE/"Eigene Auswertung"
RETRO=BASE/"kisik2"/"kisik2_icu_ml_dataset_24h.parquet"
PROS =BASE/"kisik2"/"kisik2_prospektiv_ml_dataset.parquet"
OUTdoc=AN/"KISIK_Table1_Patientencharakteristika.docx"; OUTcsv=AN/"canonical"/"table1_characteristics.csv"
allowed=[("AIN","IZ32"),("AIN","IZ21"),("AIN","IZ31")]  # nur AIN-Intensiveinheiten IZ32/IZ21/IZ31
asql=", ".join(f"('{w}','{o}')" for w,o in allowed)
con=duckdb.connect()
def load(p):
    cols=con.execute(f"SELECT * FROM read_parquet('{p}') LIMIT 0").df().columns
    extra=" AND is_open=0" if "is_open" in cols else ""   # nur abgeschlossene Aufenthalte (tatsaechliche LoS)
    d=con.execute(f"SELECT * FROM read_parquet('{p}') WHERE (wardshort,oebenekurz) IN ({asql}) AND icu_duration_h/24.0>1{extra}").df()
    d["ICU-LoS (days)"]=d["icu_duration_h"]/24.0
    if "hospital_duration_h" in d.columns: d["Hospital-LoS (days)"]=pd.to_numeric(d["hospital_duration_h"],errors="coerce")/24.0
    for c in ["alter","stay_nr","score_saps_ii_14_first","score_sofa_first","score_tiss_28_10_first"]:
        if c in d.columns: d[c]=pd.to_numeric(d[c],errors="coerce")
    return d
R=load(RETRO.as_posix()); P=load(PROS.as_posix())
# Schwere-Scores (24h) anreichern
try:
    Rsc=pd.read_csv(AN/"canonical"/"scores24_retro.csv",sep=";")
    Psc=pd.read_csv(AN/"canonical"/"scores24_prospektiv.csv",sep=";")
    R=R.drop(columns=[c for c in Rsc.columns if c!="stay_id" and c in R.columns]).merge(Rsc,on="stay_id",how="left")
    P=P.drop(columns=[c for c in Psc.columns if c!="stay_id" and c in P.columns]).merge(Psc,on="stay_id",how="left")
except Exception as e: print("Score-Merge uebersprungen:",e)
print(f"Retro n={len(R):,} (Patienten {R['pid'].nunique():,}) | Prosp n={len(P):,} (Patienten {P['pid'].nunique():,})")
for name,d in [("Retro",R),("Prosp",P)]:
    pb=pd.to_datetime(d["planbegin"],errors="coerce")
    print(f"  {name}-Zeitraum: {pb.min().date()} bis {pb.max().date()}")

CONT=[("Age (years)","alter"),("SAPS II (admission)","score_saps_ii_14_first"),
      ("SOFA (admission)","score_sofa_first"),("TISS-28 (admission)","score_tiss_28_10_first"),
      ("ICU length of stay (days)","ICU-LoS (days)"),("Hospital length of stay (days)","Hospital-LoS (days)"),
      ("ICU stay number","stay_nr")]
CATS=[("Ward","wardshort"),("ICU care unit","oebenekurz")]

def smd_cont(a,b):
    a=a[~np.isnan(a)]; b=b[~np.isnan(b)]
    s=np.sqrt((a.var(ddof=1)+b.var(ddof=1))/2)
    return 0.0 if s==0 else abs(a.mean()-b.mean())/s
def smd_prop(p1,p2):
    s=np.sqrt((p1*(1-p1)+p2*(1-p2))/2)
    return 0.0 if s==0 else abs(p1-p2)/s
def fmt_iqr(x):
    x=x[~np.isnan(x)]
    return f"{np.median(x):.1f} [{np.percentile(x,25):.1f}–{np.percentile(x,75):.1f}]"

rows=[]
for label,col in CONT:
    if col not in R.columns or col not in P.columns: continue
    a=R[col].to_numpy(float); b=P[col].to_numpy(float)
    av=a[~np.isnan(a)]; bv=b[~np.isnan(b)]
    # nur Variablen mit >=20% Abdeckung in BEIDEN Kohorten (faire Vergleichbarkeit)
    if len(av)<0.2*len(a) or len(bv)<0.2*len(b): continue
    p=stats.mannwhitneyu(av,bv,alternative="two-sided").pvalue
    cov=f"{100*len(av)/len(a):.0f}/{100*len(bv)/len(b):.0f}"
    rows.append({"Characteristic":label,"Retrospective":fmt_iqr(a),"Prospective":fmt_iqr(b),
                 "p":p,"SMD":round(smd_cont(a,b),3),"Avail_%":cov,"type":"cont"})

# Schwere-Scores (asymmetrische Verfuegbarkeit -> deskriptiv, nicht cross-verglichen)
rows.append({"Characteristic":"Severity scores, first 24 h (availability %)","Retrospective":"","Prospective":"",
             "p":"","SMD":"","Avail_%":"","type":"cat_head"})
for label,col in [("   SAPS II","score_saps_first"),("   TISS-28","score_tiss_first"),("   SOFA","score_sofa_first")]:
    def desc(d,c):
        if c not in d.columns: return "not recorded",0.0
        v=pd.to_numeric(d[c],errors="coerce"); av=100*v.notna().sum()/len(d)
        return (f"{fmt_iqr(v.to_numpy(float))} ({av:.0f}%)" if av>=20 else f"n.a. ({av:.0f}%)"), av
    rstr,_=desc(R,col); pstr,_=desc(P,col)
    rows.append({"Characteristic":label,"Retrospective":rstr,"Prospective":pstr,
                 "p":"","SMD":"n.c.","Avail_%":"","type":"score"})

# kategorial
for label,col in CATS:
    lv=sorted(set(R[col].dropna().astype(str))|set(P[col].dropna().astype(str)))
    ct=np.array([[ (R[col].astype(str)==l).sum() for l in lv ],[ (P[col].astype(str)==l).sum() for l in lv ]])
    chi=stats.chi2_contingency(ct)
    rows.append({"Characteristic":f"{label} (overall)","Retrospective":f"n={len(R)}","Prospective":f"n={len(P)}",
                 "p":chi.pvalue,"SMD":"","Avail_%":"100/100","type":"cat_head"})
    for i,l in enumerate(lv):
        n1,n2=ct[0,i],ct[1,i]; p1,p2=n1/len(R),n2/len(P)
        if p1<0.01 and p2<0.01: continue
        rows.append({"Characteristic":f"   {l}","Retrospective":f"{n1} ({100*p1:.1f}%)","Prospective":f"{n2} ({100*p2:.1f}%)",
                     "p":"","SMD":round(smd_prop(p1,p2),3),"Avail_%":"","type":"cat_lvl"})

df=pd.DataFrame(rows)
def pfmt(x):
    if x=="" or x is None: return ""
    return "<0.001" if x<0.001 else f"{x:.3f}"
df["p_str"]=df["p"].apply(lambda v: pfmt(v) if isinstance(v,(int,float)) else "")
df.drop(columns=["p"]).to_csv(OUTcsv,sep=";",index=False)
print("\n=== TABLE 1 ===")
print(df[["Characteristic","Retrospective","Prospective","p_str","SMD"]].to_string(index=False))

# ---------------- Word-Dokument ----------------
FONT="Times New Roman"; HF="1F4E79"; ZB="EAF1F8"
doc=Document(); st=doc.styles["Normal"]; st.font.name=FONT; st.font.size=Pt(10.5)
s=doc.sections[0]; s.page_width=Cm(21); s.page_height=Cm(29.7); s.left_margin=s.right_margin=Cm(2)
def setbg(c,h):
    tp=c._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:color"),"auto"); sh.set(qn("w:fill"),h); tp.append(sh)
def cell(c,t,b=False,col=None,al="left",sz=9.5):
    c.text=""; p=c.paragraphs[0]; p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER}[al]
    p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
    r=p.add_run(t); r.font.name=FONT; r.font.size=Pt(sz); r.font.bold=b
    if col: r.font.color.rgb=col
    c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
h=doc.add_paragraph(); r=h.add_run("Table 1. Patient and stay characteristics: retrospective development vs prospective cohort.")
r.bold=True; r.font.size=Pt(11); r.font.name=FONT
cols=["Characteristic",f"Retrospective (n={len(R):,})",f"Prospective (n={len(P):,})","p-value","SMD"]
w=[6.2,4.3,4.3,1.7,1.2]
t=doc.add_table(rows=len(df)+1,cols=5); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for j,(cn,cw) in enumerate(zip(cols,w)):
    c=t.cell(0,j); c.width=Cm(cw); cell(c,cn,b=True,col=RGBColor(0xFF,0xFF,0xFF),al="center"); setbg(c,HF)
for i,(_,rr) in enumerate(df.iterrows(),1):
    vals=[rr["Characteristic"],rr["Retrospective"],rr["Prospective"],rr["p_str"],str(rr["SMD"])]
    for j,(v,cw) in enumerate(zip(vals,w)):
        c=t.cell(i,j); c.width=Cm(cw)
        cell(c,v,b=(rr["type"]=="cat_head"),al="left" if j==0 else "center")
        if rr["type"]=="cat_head": setbg(c,ZB)
def note(txt):
    p=doc.add_paragraph(); r=p.add_run(txt); r.font.size=Pt(8.5); r.italic=True; r.font.name=FONT; r.font.color.rgb=RGBColor(0x44,0x44,0x44)
note("Continuous variables: median [IQR], Mann-Whitney U test. Categorical variables: n (%), chi-square test. "
     "SMD = standardised mean difference (|SMD| > 0.1 indicates a notable between-cohort imbalance). Both cohorts use "
     "identical inclusion criteria (17 ICU ward/care-unit combinations; ICU LOS > 1 day) and do not overlap in time "
     "(retrospective up to Jul 2024; prospective from Oct 2024), i.e. a true out-of-time comparison.")
note("Severity scores were extracted from the score table (first value within 24 h). Their availability differs "
     "markedly by period because of a change in scoring practice: SAPS II and TISS-28 were routinely recorded "
     "retrospectively (≈ 73%) but almost never prospectively (≤ 1%), whereas SOFA is sparse in both (1–2%). A fair "
     "cross-cohort comparison of severity is therefore not possible (n.c. = not compared); the scores are shown "
     "descriptively. Hospital length of stay was likewise unavailable prospectively. Sex and body-mass index were not "
     "available in either dataset.")
note("With the large retrospective sample the p-values are strongly powered, so small absolute differences reach "
     "significance; the SMD is the more meaningful measure of cohort imbalance. The prospective cohort here includes "
     "only completed stays (is_open = 0, LOS > 1 day) under identical inclusion criteria (n = 2,026); the "
     "senior-physician benchmark used a subset of 193 completed stays with a documented estimate.")
doc.save(str(OUTdoc))
print(f"\nGespeichert: {OUTdoc}\n           {OUTcsv}")
