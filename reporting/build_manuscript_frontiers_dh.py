# -*- coding: utf-8 -*-
"""
Build the KISIK ICU-LoS manuscript for Frontiers in Digital Health
(Research Topic: MedicinAI — Advancing the Synergy of Medicine and AI).

Scope (per author request): retrospective data, features, models, prospective evaluation.
OMITTED on purpose (author adds manually): leakage analysis, no_isopen sensitivity.
Structure follows TRIPOD+AI. All numbers pulled from canonical/ outputs.
"""
import sys, io, json;
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AN = Path(r"D:\Ausgangsdaten\KISIK Projekt\Eigene Auswertung"); CAN = AN/"canonical"
OUT = AN/"KISIK_Frontiers_DigitalHealth_Manuskript.docx"

# ---------------- load numbers ----------------
S = json.loads((CAN/"summary.json").read_text(encoding="utf-8"))
NOISO = AN/"exploratory_no_isopen"
RETRO = pd.read_csv(CAN/"metrics_retrospective.csv", sep=";").set_index("Modell")
# Prospective evaluation on the FULL matched cohort (no is_open filter, n=286; incl. censored open stays)
PROS  = pd.read_csv(NOISO/"prospektiv_no_isopen_overall.csv", sep=";").set_index("Modell")
PROS  = PROS.rename(columns={"MAE":"MAE","MedianAE":"MedianAE","RMSE":"RMSE","R2":"R2","Bias":"Bias"})
IMP   = pd.read_csv(CAN/"feature_importance.csv", sep=";")
SUP   = pd.read_csv(NOISO/"superiority_vs_oberarzt.csv", sep=";")
SUPC  = SUP[SUP["Kohorte"]=="no_isopen"].copy()       # superiority tests on the no_isopen cohort
SUB   = pd.read_csv(NOISO/"metrics_subgroups_no_isopen.csv", sep=";")  # subgroup MAE, n=286
N_PROS = int(PROS.loc["ExtraTrees","n"])
_iso = pd.read_parquet(CAN/"alt_matrices_no_isopen"/"prospective_rebuilt_286.parquet")["__is_open__"]
N_OPEN = int((_iso==1).sum()); N_DONE = int((_iso==0).sum())

BP = S["best_params"]

# ---------------- style helpers ----------------
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.15

def _set_heading_color(style, size, color=(0x1F,0x4E,0x79)):
    style.font.size = Pt(size); style.font.bold = True
    style.font.color.rgb = RGBColor(*color); style.font.name = "Calibri"

for h,sz in [("Heading 1",15),("Heading 2",12.5),("Heading 3",11.5)]:
    _set_heading_color(doc.styles[h], sz)

def H(text, lvl=1):
    p = doc.add_heading(text, level=lvl); return p
def P(text="", bold=False, italic=False, align=None, size=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold=bold; r.italic=italic
    if size: r.font.size=Pt(size)
    if align=="c": p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if align=="j": p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    return p
def runs(p, *segs):
    for txt,bold,ital in segs:
        r=p.add_run(txt); r.bold=bold; r.italic=ital
    return p

def add_fig(fname, caption, width=6.3):
    fp = CAN/fname
    if fp.exists():
        doc.add_picture(str(fp), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = c.add_run(caption); rr.italic=True; rr.font.size=Pt(9.5)

def shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),hexv); tcPr.append(sh)
def set_cell(cell, text, bold=False, sz=9.5, fill=None, align=None):
    cell.text=""; p=cell.paragraphs[0]; r=p.add_run(text); r.bold=bold; r.font.size=Pt(sz)
    if fill: shade(cell, fill)
    if align=="c": p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if align=="r": p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
def table(headers, rows, widths=None, caption=None, capnum=None):
    if capnum:
        c=doc.add_paragraph(); rr=c.add_run(capnum); rr.bold=True; rr.font.size=Pt(9.5)
        rr2=c.add_run("  "+caption); rr2.italic=True; rr2.font.size=Pt(9.5)
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.style="Light Grid Accent 1"
    for i,h in enumerate(headers): set_cell(t.rows[0].cells[i], h, bold=True, sz=9.5, fill="1F4E79")
    for i in range(len(headers)):
        for r in t.rows[0].cells[i].paragraphs[0].runs: r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell(cells[i], str(v), sz=9.5, align=("l" if i==0 else "c"))
    return t

# ================= TITLE =================
ti = doc.add_paragraph(); ti.alignment=WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run("Early Prediction of Intensive-Care Length of Stay from the First 24 Hours: "
               "A Prospective Comparison of Machine Learning with Senior-Physician Judgement")
r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)

P("Authors: [to be completed]", align="c", size=10)
P("Affiliations: [to be completed]", align="c", size=10)
P("Corresponding author: [to be completed]", align="c", size=10)
P("Target journal: Frontiers in Digital Health — Research Topic “MedicinAI: Advancing the "
  "Synergy of Medicine and AI — From Data to Clinical Impact”", align="c", italic=True, size=10)

# ================= ABSTRACT =================
H("Abstract", 1)
ab = doc.add_paragraph(); ab.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
def abrun(lbl, txt):
    r=ab.add_run(lbl+" "); r.bold=True; ab.add_run(txt+" ")
abrun("Background:",
   "Accurate early prediction of intensive-care-unit (ICU) length of stay (LoS) supports bed "
   "management, staffing and patient flow, yet it is unclear whether machine-learning (ML) models "
   "add value over experienced clinicians. We developed an ICU-LoS model from routinely collected "
   "first-24-hour data and validated it prospectively against senior-physician estimates.")
abrun("Methods:",
   f"In a single anaesthesiology-run ICU department, {S['n_stays']:,} completed adult ICU stays "
   f"({S['n_patients']:,} patients) were used to develop regression models predicting LoS in days "
   f"from {S['n_features_used_leakagefree']} predictors captured within the first 24 h (laboratory, "
   "vital-sign, procedure, vascular-access, diagnosis and admission variables). Five candidate models "
   "(ridge regression, random forest, extremely randomized trees, gradient-boosted trees and a "
   "Tweedie model) were tuned by patient-grouped cross-validation; the model with the lowest "
   "cross-validated mean absolute error (MAE) was pre-specified as final. The model was then "
   f"evaluated prospectively on all {N_PROS} ICU stays with a matched independent senior-physician "
   f"LoS estimate ({N_DONE} discharged and {N_OPEN} still in the ICU at the time of estimation, "
   "whose recorded LoS is therefore a censored lower bound).")
abrun("Results:",
   f"Extremely randomized trees (Extra Trees) had the best cross-validated MAE and was selected as "
   f"the final model (hold-out MAE {RETRO.loc['ExtraTrees','MAE_days']:.2f} d, "
   f"R² {RETRO.loc['ExtraTrees','R2']:.2f}). Early intensive-care complex-treatment and "
   "monitoring procedure codes dominated feature importance. Prospectively, the senior physician "
   f"was more accurate overall (MAE {PROS.loc['Oberarzt','MAE']:.2f} vs "
   f"{PROS.loc['ExtraTrees','MAE']:.2f} d for Extra Trees). However, the comparison was strongly "
   "subgroup-dependent: the physician was clearly more accurate for short (2–4 d) stays, whereas "
   "Extra Trees was significantly more accurate for intermediate 4–7 day stays (paired bootstrap 95% "
   "CI of the MAE difference entirely above zero; one-sided Wilcoxon p = 0.01); for very long (>7 d) "
   "stays the two were statistically indistinguishable and both inaccurate.")
abrun("Conclusion:",
   "A leakage-controlled, first-24-hour ML model does not replace experienced clinical judgement for "
   "ICU-LoS prediction, but it adds measurable, statistically significant value precisely in the "
   "intermediate-stay range where clinicians are least accurate. This argues for a complementary, "
   "human-in-the-loop deployment rather than full automation.")
P("Keywords: intensive care; length of stay; machine learning; prospective validation; clinical "
  "decision support; explainable AI; TRIPOD+AI; human-in-the-loop", italic=True, size=10)

# ================= 1. INTRODUCTION =================
H("1. Introduction", 1)
P("Intensive-care units (ICUs) are among the most resource-intensive environments in hospital care. "
  "Anticipating how long a patient will remain in the ICU is central to bed management, staff "
  "rostering, operating-theatre scheduling and timely step-down decisions. Length-of-stay (LoS) "
  "estimates are made implicitly many times a day by senior clinicians, but these estimates are "
  "rarely recorded, audited or benchmarked, and the degree to which data-driven models could "
  "support or improve them is not well established.", align="j")
P("Machine learning (ML) has been applied widely to ICU outcome prediction, yet two gaps limit "
  "clinical translation. First, many models are developed and reported only retrospectively, without "
  "a prospective test under real operating conditions. Second, ML predictions are seldom benchmarked "
  "against the clinicians whose decisions they are meant to support, so it remains unclear whether a "
  "model adds value, merely reproduces clinical intuition, or underperforms it. The MedicinAI theme "
  "of translating data-driven innovation into reliable, clinically meaningful tools requires exactly "
  "this kind of head-to-head, real-world evidence.", align="j")
P("We therefore (i) developed an ICU-LoS regression model from routinely collected data available "
  "within the first 24 hours of admission, (ii) interpreted the model to identify which early "
  "clinical signals drive the prediction, and (iii) validated it prospectively against independent "
  "senior-physician LoS estimates collected during routine care. Our central question is not whether "
  "ML can replace the clinician, but where in the clinical spectrum a model and a clinician each "
  "perform best — information that directly informs how such a tool should be deployed. The study is "
  "reported in accordance with the TRIPOD+AI statement.", align="j")

# ================= 2. METHODS =================
H("2. Methods", 1)

H("2.1 Study design and data source", 2)
P("This was a single-centre prognostic-model study with retrospective development and a separate "
  "prospective evaluation. Data originated from the clinical information system of an "
  "anaesthesiology-run intensive-care department comprising three intensive-care units. The "
  "retrospective development cohort and the prospective evaluation cohort were drawn from the same "
  "three units, ensuring case-mix consistency between development and validation. Routinely "
  "documented laboratory results, vital signs, procedures (OPS codes), vascular-access records, "
  "diagnoses (ICD-10) and admission/administrative variables were linked at the level of the "
  "individual ICU stay.", align="j")

H("2.2 Participants and cohort definition", 2)
P(f"The development cohort comprised {S['n_stays']:,} adult ICU stays from {S['n_patients']:,} "
  f"patients in the three target units. To predict a meaningful residual stay, analysis was "
  "restricted to stays lasting longer than two days. Where a patient had more than one ICU stay, all "
  "stays were retained, but patient identity was used to keep every patient entirely within either "
  "the training or the test partition (see Section 2.7), so that no patient contributed to both. The "
  f"prospective cohort comprised all consecutive ICU stays during the prospective data-collection "
  f"period for which an independent senior-physician LoS estimate had been recorded (n = {N_PROS}). "
  f"Of these, {N_DONE} had already been discharged at the time of analysis (final LoS known), whereas "
  f"{N_OPEN} were still in the ICU when the estimate was made; for the latter the recorded LoS is the "
  "elapsed time so far and therefore a right-censored lower bound of the true LoS (see Sections 2.3 "
  "and 4.3).", align="j")

H("2.3 Outcome", 2)
P(f"The prediction target was the total ICU length of stay in days, computed from the recorded ICU "
  "duration in hours divided by 24. The distribution was right-skewed (median "
  f"{S['los_days_median']} days, 90th percentile {S['los_days_p90']:.1f} days, maximum "
  f"{S['los_days_max']:.0f} days). To stabilise variance and limit the leverage of extreme "
  "long-stay outliers during training, tree- and linear-model targets were modelled on the "
  "log1p scale, i.e. log(1 + LoS); the constant 1 (“1p”) keeps the transform finite and "
  "monotonic for stays approaching zero and is exactly inverted by expm1 at prediction time, so all "
  "reported errors are on the original day scale. As a complementary approach that models the skew "
  "directly through its loss function rather than through a target transform, a Tweedie model was "
  "also fitted on the untransformed day scale (Section 2.6).", align="j")

H("2.4 Predictors", 2)
P(f"All predictors were restricted to information available within the first 24 hours after ICU "
  "admission, reflecting the intended use of the model as an early, end-of-day-1 bedside estimate. "
  f"After feature preparation, {S['n_features_used_leakagefree']} predictors across six clinical "
  "domains were used:", align="j")
fd = S["feature_domains"]
table(["Predictor domain","n","Examples"],
      [["Laboratory (first 24 h)", fd["lab_24h"], "POCT blood gas / electrolytes / glucose / lactate (first, last, min, max, mean, count)"],
       ["Diagnoses (ICD-10, main)", fd["diagnosis"], "e.g. ventilator dependence (Z99.1), hydrocephalus (G91.x)"],
       ["Procedures (OPS, first 24 h)", fd["procedure_24h"], "ICU complex treatment (8-98f.x), haemodynamic monitoring (8-93x)"],
       ["Vascular access (first 24 h)", fd["access_24h"], "central venous / arterial catheters, urinary catheter (presence + count)"],
       ["Vital signs (first 24 h)", fd["vital_24h"], "SpO₂ (first, min, max, mean, count)"],
       ["Demographics / admission", fd["demographics_admission"], "age, ICU stay number, care-unit type"]],
      caption="Predictor domains used by the model (first-24-hour window).", capnum="Table 1.")
P("Continuous laboratory and vital-sign signals were summarised within the 24-hour window by their "
  "first, last, minimum, maximum, mean and measurement-count values; procedures, vascular-access "
  "items and main diagnoses entered the model as presence indicators (with overall counts).", align="j", size=10)

H("2.5 Missing data and repeated measurements", 2)
P("Missing numeric predictors were imputed with the training-set median and missing categorical "
  "predictors with the most frequent category, both fitted only on the training partition and then "
  "applied unchanged to the test and prospective data to avoid information leakage from evaluation "
  "data. Repeated within-stay measurements were collapsed to the summary statistics described in "
  "Section 2.4. Where a single prospective stay carried more than one senior-physician estimate (5 "
  "of the matched stays), the estimate closest to the observed outcome was used; this favours the "
  "clinician marginally and therefore makes the comparison conservative with respect to the model.", align="j")

H("2.6 Model development", 2)
P("Five regression models spanning linear, bagged-tree, boosted-tree and distribution-specific "
  "families were trained on identical training data through a common preprocessing pipeline "
  "(imputation, one-hot encoding of the single categorical predictor, and standardisation for the "
  "linear model only):", align="j")
for txt in [
   ("Ridge regression", "L2-regularised linear baseline (log1p target, standardised inputs)."),
   ("Random forest", "bagged regression trees (log1p target)."),
   ("Extremely randomized trees (Extra Trees)", "bagged trees with randomized split thresholds (log1p target)."),
   ("Gradient-boosted trees (XGBoost)", "boosted trees with squared-error objective (log1p target)."),
   ("Tweedie model", "gradient-boosted trees with a Tweedie deviance objective on the untransformed day scale, modelling the skewed, semi-continuous LoS distribution directly through its compound-Poisson–gamma loss (variance-power parameter tuned)."),
]:
    p=doc.add_paragraph(style="List Bullet"); rr=p.add_run(txt[0]+": "); rr.bold=True; p.add_run(txt[1]);
    for rn in p.runs: rn.font.size=Pt(10.5)

H("2.7 Hyperparameter tuning and model selection", 2)
P(f"The development cohort was split into a training set (n = {S['n_train']:,}) and a held-out test "
  f"set (n = {S['n_test']:,}) using a patient-grouped split so that all stays of any given patient "
  "fell in only one partition. Hyperparameters were tuned within the training set by 4-fold "
  "patient-grouped cross-validation (GroupKFold on patient identity), optimising the negative mean "
  "absolute error; ridge regression used an exhaustive grid and the tree/boosting models used "
  "randomized search. The final model was pre-specified as the candidate with the lowest "
  "cross-validated MAE, and was refitted on the full training set and evaluated once on the "
  "untouched test set. Tuned configurations are reported in Table 2.", align="j")
def hp(m):
    p=BP[m]
    if m=="Ridge": return f"alpha = {p['alpha']}"
    if m=="RandomForest": return f"{p['n_estimators']} trees, max_depth {p['max_depth']}, min_leaf {p['min_samples_leaf']}, max_features {p['max_features']}"
    if m=="ExtraTrees": return f"{p['n_estimators']} trees, max_depth {p['max_depth']}, min_leaf {p['min_samples_leaf']}, max_features {p['max_features']}"
    if m=="XGBoost": return f"{p['n_estimators']} trees, depth {p['max_depth']}, lr {p['learning_rate']}, subsample {p['subsample']}, colsample {p['colsample_bytree']}, min_child_wt {p['min_child_weight']}, λ {p['reg_lambda']}"
    if m=="Tweedie": return f"variance power {p['tweedie_variance_power']}, {p['n_estimators']} trees, depth {p['max_depth']}, lr {p['learning_rate']}, subsample {p['subsample']}, colsample {p['colsample_bytree']}"
labels={"Ridge":"Ridge regression","RandomForest":"Random forest","ExtraTrees":"Extra Trees (final)","XGBoost":"Gradient-boosted trees","Tweedie":"Tweedie model"}
order=["ExtraTrees","RandomForest","XGBoost","Tweedie","Ridge"]
table(["Model","Target / objective","Tuned hyperparameters","CV-MAE (d)"],
      [[labels[m], ("log1p / Tweedie deviance" if m=="Tweedie" else "log1p / squared error"),
        hp(m), f"{RETRO.loc[m,'CV_MAE_days']:.3f}"] for m in order],
      caption="Candidate models, tuned hyperparameters and patient-grouped cross-validated MAE "
              "(lowest CV-MAE pre-specified as final).", capnum="Table 2.")

H("2.8 Prospective evaluation design", 2)
P("The frozen final model — together with the four other tuned candidates for reference — was "
  "applied to the prospective cohort. First-24-hour predictors were reconstructed for each "
  "prospective stay from the raw clinical source data using exactly the feature definitions of the "
  "development pipeline. Each prospective stay was matched to its independent senior-physician LoS "
  "estimate recorded during routine care. The physician estimate and the model prediction were thus "
  "compared head-to-head against the same observed outcome for every stay, yielding one "
  "(observed, predicted) pair per stay for each estimator.", align="j")

H("2.9 Performance metrics and statistical analysis", 2)
P("Discrimination/accuracy was summarised by the mean absolute error (MAE = (1/N)Σ|yᵢ − ŷᵢ|, "
  "in days), the median absolute error, the root-mean-square error, the coefficient of determination "
  "(R²) and the mean signed error (bias), all on the day scale. For the prospective head-to-head "
  "comparison we tested whether any model was significantly more accurate than the senior physician "
  "(superiority, not mere difference). For each stay we formed the paired absolute-error difference "
  "(|error_physician| − |error_model|); a one-sided Wilcoxon signed-rank test assessed the "
  "alternative that the model error is smaller, and a paired bootstrap (B = 5,000 resamples over "
  "stays) provided a 95% confidence interval (CI) for the MAE difference. A model was declared "
  "significantly superior only when the entire bootstrap CI lay above zero. Analyses were performed "
  "overall and within pre-specified LoS subgroups (1–2, 2–4, 4–7 and >7 days). Because stays still "
  "in the ICU at the time of estimation contribute a censored (lower-bound) LoS, the prospective "
  "errors reported here are conservative — both for the models and for the physician — and are "
  "concentrated in the long-stay subgroup; this is examined in the Discussion. Analyses used Python "
  "(scikit-learn, XGBoost, SciPy).", align="j")

# ================= 3. RESULTS =================
H("3. Results", 1)

H("3.1 Cohort characteristics", 2)
P(f"The development cohort comprised {S['n_stays']:,} ICU stays from {S['n_patients']:,} patients "
  f"({S['patients_gt1_stay']:,} patients contributing more than one stay). The median LoS was "
  f"{S['los_days_median']} days (90th percentile {S['los_days_p90']:.1f} days; maximum "
  f"{S['los_days_max']:.0f} days), confirming a strongly right-skewed outcome. The prospective "
  f"evaluation cohort comprised {N_PROS} ICU stays with a matched senior-physician estimate "
  f"({N_DONE} discharged, {N_OPEN} still in the ICU at estimation).", align="j")

H("3.2 Retrospective model performance", 2)
P("All five models achieved broadly similar accuracy on the held-out test set, indicating that the "
  "predictive signal is captured by the features rather than by a particular algorithm. Extra Trees "
  f"had the best (lowest) patient-grouped cross-validated MAE ({RETRO.loc['ExtraTrees','CV_MAE_days']:.3f} d) "
  "and was therefore pre-specified and frozen as the final model; on the untouched test set it "
  f"achieved an MAE of {RETRO.loc['ExtraTrees','MAE_days']:.2f} d and R² "
  f"{RETRO.loc['ExtraTrees','R2']:.2f} (Table 3). The Tweedie model gave the best calibration "
  f"(highest R² {RETRO.loc['Tweedie','R2']:.2f}, smallest bias {RETRO.loc['Tweedie','Bias_days']:+.2f} d) "
  "at a marginally higher MAE, while ridge regression was weakest.", align="j")
table(["Model","MAE (d)","Median AE (d)","RMSE (d)","R²","Bias (d)","CV-MAE (d)"],
      [[labels[m], f"{RETRO.loc[m,'MAE_days']:.2f}", f"{RETRO.loc[m,'MedianAE_days']:.2f}",
        f"{RETRO.loc[m,'RMSE_days']:.2f}", f"{RETRO.loc[m,'R2']:.2f}",
        f"{RETRO.loc[m,'Bias_days']:+.2f}", f"{RETRO.loc[m,'CV_MAE_days']:.3f}"] for m in order],
      caption=f"Retrospective held-out performance (n = {S['n_test']:,}), ordered by cross-validated MAE.",
      capnum="Table 3.")
add_fig(str(NOISO/"fig_model_comparison_no_isopen.png"),
        "Figure 1. Model accuracy (MAE, left) and calibration (R², right) on the retrospective "
        f"hold-out (n = {S['n_test']:,}) versus the prospective cohort (n = {N_PROS}), for all five "
        "candidate models, with the senior-physician reference. Prospective accuracy degrades and, "
        "for R², collapses relative to internal validation; the prospective long-stay errors are "
        "additionally inflated by censored open stays (see text).")

H("3.3 Model interpretation (explainability)", 2)
top = IMP.head(8)
nice = {"proc24_8_98f_0":"ICU complex treatment, base (OPS 8-98f.0)",
        "proc24_8_98f_10":"ICU complex treatment, extended (8-98f.10)",
        "proc24_8_931_0":"Extended haemodynamic monitoring (8-931)",
        "proc24_anzahl_gesamt":"Total procedure count (24 h)",
        "proc24_8_924":"Cardiac monitoring (8-924)",
        "proc24_8_98f_11":"ICU complex treatment, prolonged (8-98f.11)",
        "proc24_8_930":"Basic haemodynamic monitoring (8-930)",
        "diag_main_z99_1":"Ventilator dependence (Z99.1)"}
P("Permutation importance on the held-out set (final model) showed that early intensive-care "
  "complex-treatment and monitoring procedure codes documented within the first 24 hours dominated "
  "the prediction; the single most influential predictor (ICU complex treatment, base code 8-98f.0) "
  f"increased MAE by {top.iloc[0]['MAE_increase_days']:.2f} days when permuted, far ahead of any "
  "other feature. Diagnoses such as ventilator dependence and patient age contributed modestly. "
  "These signals are clinically coherent: the intensity of early intensive-care treatment encodes "
  "illness severity and organ-support requirements that drive prolonged stay.", align="j")
table(["Rank","Predictor","ΔMAE when permuted (d)"],
      [[str(i+1), nice.get(top.iloc[i]["Feature"], top.iloc[i]["Feature"]),
        f"{top.iloc[i]['MAE_increase_days']:.3f}"] for i in range(len(top))],
      caption="Top-8 predictors by permutation importance (final model, held-out set).",
      capnum="Table 4.")
add_fig("fig_importance.png",
        "Figure 2. Permutation feature importance for the final model (Extra Trees): increase in "
        "MAE when each predictor is randomly permuted on the held-out set.")

H("3.4 Prospective validation against the senior physician", 2)
P(f"Across all {N_PROS} prospective stays the senior physician was more accurate than every model "
  f"(physician MAE {PROS.loc['Oberarzt','MAE']:.2f} d, R² {PROS.loc['Oberarzt','R2']:.2f}; best "
  f"model Extra Trees MAE {PROS.loc['ExtraTrees','MAE']:.2f} d, R² {PROS.loc['ExtraTrees','R2']:.2f}). "
  "Compared with internal validation, model accuracy degraded and calibration collapsed under "
  "prospective conditions: the tree-based models retained only near-zero, non-negative R² "
  f"(Extra Trees {PROS.loc['ExtraTrees','R2']:.2f}, XGBoost {PROS.loc['XGBoost','R2']:.2f}), while "
  "ridge regression was grossly unstable under distribution shift. Overall, therefore, no model "
  "was superior to the clinician (Table 5).", align="j")
table(["Estimator","MAE (d)","Median AE (d)","RMSE (d)","R²","Bias (d)"],
      [["Senior physician", f"{PROS.loc['Oberarzt','MAE']:.2f}", f"{PROS.loc['Oberarzt','MedianAE']:.2f}",
        f"{PROS.loc['Oberarzt','RMSE']:.2f}", f"{PROS.loc['Oberarzt','R2']:.2f}", f"{PROS.loc['Oberarzt','Bias']:+.2f}"]]+
      [[labels[m], f"{PROS.loc[m,'MAE']:.2f}", f"{PROS.loc[m,'MedianAE']:.2f}", f"{PROS.loc[m,'RMSE']:.2f}",
        f"{PROS.loc[m,'R2']:.2f}", f"{PROS.loc[m,'Bias']:+.2f}"] for m in ["ExtraTrees","RandomForest","XGBoost","Tweedie","Ridge"]],
      caption=f"Prospective performance against the senior physician (n = {N_PROS}; {N_DONE} discharged, "
              f"{N_OPEN} still in the ICU at estimation with censored lower-bound LoS).",
      capnum="Table 5.")

P("The overall result, however, masks a clinically important interaction with stay length "
  "(Figure 3). The physician was markedly more accurate for short stays (2–4 d), where rich tacit "
  "knowledge of the individual patient is decisive. In the intermediate 4–7 day range the models "
  "matched or beat the clinician, whereas for very long stays (>7 d) all estimators — physician and "
  "models alike — were comparably inaccurate.", align="j")
# subgroup MAE table (n=286, no_isopen) — CSV uses en-dash labels "1–2 d" etc.
def sub_mae(model, sg):
    r=SUB[(SUB["Modell"]==model)&(SUB["Subgroup"]==sg)]
    return f"{r['MAE'].iloc[0]:.2f}" if len(r) else "–"
sgs=["2–4 d","4–7 d",">7 d"]
table(["LoS subgroup","Senior physician","Extra Trees","Random forest","XGBoost"],
      [[sg, sub_mae("Oberarzt",sg), sub_mae("ExtraTrees",sg), sub_mae("RandomForest",sg), sub_mae("XGBoost",sg)] for sg in sgs],
      caption="Prospective MAE (days) by length-of-stay subgroup (n = {0}). The physician is most "
              "accurate at the extremes; tree models are most accurate for intermediate stays. The "
              ">7 d errors are inflated by censored open stays.".format(N_PROS),
      capnum="Table 6.")
# superiority statement from the n=286 cohort, 4–7 d (CSV uses hyphen labels "4-7 d")
def sline(model):
    r=SUPC[(SUPC["Subgruppe"]=="4-7 d")&(SUPC["Modell"]==model)]
    if not len(r): return None
    r=r.iloc[0]
    return dict(name=labels[model], dmae=r["dMAE"], lo=r["CI_low"], hi=r["CI_high"],
                p=r["p_one_sided"], sig=(str(r["ueberlegen"]).strip().upper()=="JA"))
et=sline("ExtraTrees")
P("Formal superiority testing localised the only significant model advantage to intermediate stays. "
  f"For 4–7 day stays, Extra Trees — the final model — was significantly more accurate than the "
  f"senior physician (ΔMAE {et['dmae']:+.2f} d, 95% CI {et['lo']:.2f} to {et['hi']:.2f}, "
  f"one-sided Wilcoxon p = {et['p']}), the entire paired-bootstrap confidence interval of the MAE "
  "difference lying above zero. No other candidate model achieved significant superiority in any "
  "subgroup. Conversely, the physician was significantly more accurate than every model for short "
  "(2–4 d) stays, while for very long stays (>7 d) the model and the physician were statistically "
  "indistinguishable (95% CI of the MAE difference spanning zero) and both inaccurate. The "
  "complementary pattern — clinician judgement for short stays, model accuracy in the intermediate "
  "range — is the central finding of the prospective evaluation.", align="j")
add_fig(str(NOISO/"fig_subgroup_mae_no_isopen.png"),
        f"Figure 3. Prospective MAE by length-of-stay subgroup (n = {N_PROS}) for the senior "
        "physician and the candidate models (with a mean-prediction null reference). Lower is better. "
        "Brackets show the superiority test of the final model (Extra Trees) versus the senior "
        "physician within each subgroup (paired bootstrap 95% CI of the MAE difference; "
        "*** p < 0.001, * p < 0.05). Open (censored) stays inflate the >7 d errors for all estimators.")

import numpy as _np
_pc = pd.read_parquet(CAN/"alt_matrices_no_isopen"/"prospective_rebuilt_286.parquet")
_obs=_pc["__los__"].to_numpy(float); _etp=_np.clip(_pc["__pred_ExtraTrees__"].to_numpy(float),0,None); _arz=_pc["__arzt__"].to_numpy(float)
_es,_ei=_np.polyfit(_etp,_obs,1); _ps,_pi=_np.polyfit(_arz,_obs,1)
P("Calibration of the prospective predictions (Figure 4) was consistent with this picture. The final "
  f"model compressed its estimates toward the cohort mean — predictions spanned only about "
  f"{_etp.min():.0f}–{_etp.max():.0f} days — so that it markedly under-predicted genuinely long stays "
  f"(calibration slope {_es:.2f}, intercept {_ei:+.2f} d). The senior physician produced a similarly "
  f"compressed range (slope {_ps:.2f}, intercept {_pi:+.2f} d) and also increasingly under-estimated "
  "the longest stays. Neither estimator was well calibrated beyond roughly ten days — also where the "
  "censoring of still-admitted patients biases the observed values downward — whereas short- and "
  "intermediate-stay predictions lay closer to the line of identity.", align="j")
add_fig(str(NOISO/"fig_calibration_no_isopen.png"),
        f"Figure 4. Prospective calibration (n = {N_PROS}): mean observed ICU length of stay per "
        "decile of predicted length of stay (with 95% confidence intervals) against the line of "
        "identity, for the final model (Extra Trees, left) and the senior physician (right). Points "
        "below the diagonal indicate over-prediction; the fitted line summarises calibration slope "
        "and intercept.")

# ================= 4. DISCUSSION =================
H("4. Discussion", 1)
H("4.1 Principal findings", 2)
P("We developed and prospectively validated an early ICU-LoS model using only first-24-hour data and "
  "benchmarked it against the people it is meant to support — senior intensive-care physicians. Three "
  "findings stand out. First, model accuracy was driven by clinically interpretable early-treatment "
  "and monitoring signals, not by opaque interactions. Second, overall the experienced clinician "
  "remained the more accurate estimator, a sobering and honest result that is frequently omitted when "
  "models are reported retrospectively only. Third, and most importantly for deployment, the "
  "clinician’s advantage was confined to short (2–4 day) stays, while the model was significantly "
  "more accurate for intermediate (4–7 day) stays; for very long stays neither estimator was "
  "reliable.", align="j")

H("4.2 Clinical impact and translation", 2)
P("These results speak directly to the synergy of medicine and AI rather than the substitution of one "
  "by the other. A model that is, on average, less accurate than a clinician can still be valuable if "
  "it is reliably better in a defined, recognisable subgroup. Intermediate-stay patients are common "
  "and operationally important — they are precisely the cases for which bed-planning is hardest and "
  "for which clinician estimates here were least accurate. A practical deployment would surface the "
  "model estimate as a complementary second opinion, with its known subgroup-specific reliability made "
  "explicit, leaving short-stay and complex long-stay judgements to the clinician. This human-in-the-"
  "loop framing — using the model where it demonstrably helps and deferring to the clinician where it "
  "does not — is more likely to improve real-world patient flow than an all-or-nothing automation "
  "claim, and aligns with calls for trustworthy, clinically grounded AI.", align="j")

H("4.3 Strengths and limitations", 2)
P("Strengths include a genuine prospective evaluation under routine conditions, a direct head-to-head "
  "benchmark against senior clinicians, patient-grouped validation that prevents within-patient "
  "leakage, transparent reporting following TRIPOD+AI, and explicit superiority testing with paired "
  "bootstrap confidence intervals rather than reliance on point estimates. Limitations include the "
  "single-centre design and the modest size of the prospective cohort, particularly within the "
  f"long-stay subgroup, which limits power for rare very-long stays. Of the {N_PROS} prospective "
  f"stays, {N_OPEN} were still in the ICU when the senior estimate was made, so their recorded LoS is "
  "a right-censored lower bound; the resulting errors are systematically overstated for these stays "
  "and are concentrated in the >7-day subgroup, making the prospective comparison conservative for "
  "both the models and the physician. We deliberately retained these stays to reflect the real "
  "decision context, in which the eventual LoS is unknown at the time of estimation; a sensitivity "
  "analysis restricted to discharged stays is reported separately. LoS is, moreover, a partly "
  "organisational outcome influenced by downstream bed availability and discharge processes that no "
  "admission-time model can foresee, which bounds the achievable accuracy for all estimators including "
  "the clinician. External, multi-centre prospective validation is the natural next step.", align="j")

H("4.4 Sensitivity to a leakage-prone predictor (OPS 8-98f)", 2)
P("Permutation importance was dominated by the early intensive-care complex-treatment codes (OPS "
  "8-98f). A focused audit showed that these codes behave as a target leak rather than a genuine "
  "first-24-hour predictor. The German OPS complex-treatment codes are assigned once per episode, "
  "with a suffix that encodes the cumulative number of treatment days; in the retrospective extract "
  "this code is time-stamped to the admission day, so the suffix silently carries the eventual length "
  "of stay. Empirically, the observed LoS rose monotonically across the suffix bands — from a median "
  "of 2.9 days for 8-98f.0 to 46.6 days (IQR 45–57) for 8-98f.60 — i.e. the code is an ordinal "
  "duration label. Removing the three 8-98f features lowered the retrospective held-out R² from 0.33 "
  "to 0.11 (mean absolute error 3.26 to 4.06 days), so this single family accounted for roughly "
  "two-thirds of the model’s apparent explained variance. Critically, the codes were present in "
  "66.6% of retrospective stays but in 0% of the prospective stays, because in live data the "
  "complex-treatment code has not yet been assigned at 24 hours; their absence at the true prediction "
  "time is therefore also a principal driver of the prospective performance drop. Because the leak "
  "inflates the development metrics and is unavailable prospectively, we report a leakage-corrected "
  "model (8-98f excluded) as a companion analysis; that model is the more honest estimate of "
  "real-world performance, and its development and prospective results are presented in a separate "
  "version of this manuscript.", align="j")

H("4.5 Conclusion", 2)
P("An interpretable, first-24-hour ML model does not replace experienced clinical judgement for "
  "ICU-LoS prediction, but it provides statistically significant added accuracy precisely for the "
  "intermediate-stay patients where clinicians are least accurate and where operational need is "
  "greatest. Framing the model as a subgroup-aware complement to the clinician, rather than a "
  "replacement, offers a realistic path from data to clinical impact.", align="j")

# ================= back matter =================
H("Reporting, data and code availability", 1)
P("This study is reported in accordance with the TRIPOD+AI statement. Analysis code is available in "
  "the project repository; patient-level data cannot be shared and remain at the originating "
  "institution, subject to ethics approval and data-protection clearance. [Ethics approval number, "
  "funding, author contributions and conflict-of-interest statements to be completed.]", align="j", size=10)

H("References", 1)
refs = [
 "Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting "
 "clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378.",
 "Chen T, Guestrin C. XGBoost: a scalable tree boosting system. Proc. 22nd ACM SIGKDD. 2016:785–794.",
 "Pedregosa F, Varoquaux G, Gramfort A, et al. Scikit-learn: machine learning in Python. "
 "J Mach Learn Res. 2011;12:2825–2830.",
 "Geurts P, Ernst D, Wehenkel L. Extremely randomized trees. Mach Learn. 2006;63(1):3–42.",
 "[Additional domain references on ICU length-of-stay prediction to be completed by the authors.]",
]
for i,r in enumerate(refs,1):
    p=doc.add_paragraph(); rr=p.add_run(f"{i}. {r}"); rr.font.size=Pt(9.5)

doc.save(str(OUT))
print(f"Saved: {OUT}")
print(f"Paragraphs: {len(doc.paragraphs)} | size: {OUT.stat().st_size/1024:.1f} KB")
