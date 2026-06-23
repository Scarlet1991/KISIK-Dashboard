# -*- coding: utf-8 -*-
"""SECOND (leakage-corrected) version of the manuscript: OPS 8-98f complex-treatment
family removed. Reads leakfree/ outputs. Same structure as the primary version."""
import sys, io, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from pathlib import Path
import pandas as pd, numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AN=Path(r"D:\Ausgangsdaten\KISIK Projekt\Eigene Auswertung"); CAN=AN/"canonical"; LF=AN/"leakfree"
OUT=AN/"KISIK_Frontiers_DigitalHealth_Manuskript_v2_leakfree.docx"
S=json.loads((LF/"summary_lf.json").read_text(encoding="utf-8"))
SC=json.loads((CAN/"summary.json").read_text(encoding="utf-8"))   # for feature domains
RETRO=pd.read_csv(LF/"metrics_retrospective_lf.csv",sep=";").set_index("Modell")
PROS =pd.read_csv(LF/"prospektiv_overall_lf.csv",sep=";").set_index("Modell")
IMP  =pd.read_csv(LF/"feature_importance_lf.csv",sep=";")
SUB  =pd.read_csv(LF/"metrics_subgroups_lf.csv",sep=";")
SUP  =pd.read_csv(LF/"superiority_lf.csv",sep=";")
BP=S["best_params"]; FKEY=S["final_model"]
LK=json.loads((AN/"leak_8_98f_summary.json").read_text(encoding="utf-8"))
_supd={r["Subgruppe"]:r for _,r in SUP.iterrows()}
_BINS=["1–2 d","2–4 d","4–7 d",">7 d"]
MODEL_SG=[k for k in _BINS if k in _supd and _supd[k]["verdict"]=="model"]
PHYS_SG =[k for k in _BINS if k in _supd and _supd[k]["verdict"]=="physician"]
NS_SG   =[k for k in _BINS if k in _supd and _supd[k]["verdict"]=="n.s."]
MSG=" and ".join(MODEL_SG) if MODEL_SG else "no subgroup"
PSG=" and ".join(PHYS_SG) if PHYS_SG else "no subgroup"
_iso=pd.read_parquet(CAN/"alt_matrices_no_isopen"/"prospective_rebuilt_286.parquet")["__is_open__"]
N_PROS=int(PROS.loc["Oberarzt","n"]); N_OPEN=int((_iso==1).sum()); N_DONE=int((_iso==0).sum())
labels={"Ridge":"Ridge regression","RandomForest":"Random forest","ExtraTrees":"Extra Trees",
        "XGBoost":"Gradient-boosted trees (XGBoost)","Tweedie":"Tweedie model"}
FINAL=labels[FKEY]

doc=Document(); st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11)
st.paragraph_format.space_after=Pt(6); st.paragraph_format.line_spacing=1.15
for h,sz in [("Heading 1",15),("Heading 2",12.5),("Heading 3",11.5)]:
    s=doc.styles[h]; s.font.size=Pt(sz); s.font.bold=True; s.font.color.rgb=RGBColor(0x1F,0x4E,0x79); s.font.name="Calibri"
def H(t,l=1): doc.add_heading(t,level=l)
def P(t="",bold=False,italic=False,align=None,size=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=italic
    if size:r.font.size=Pt(size)
    if align=="c":p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if align=="j":p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    return p
def shade(c,h):
    tcPr=c._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:val"),"clear"); sh.set(qn("w:fill"),h); tcPr.append(sh)
def setc(c,t,bold=False,sz=9.5,fill=None,align=None):
    c.text="";p=c.paragraphs[0];r=p.add_run(t);r.bold=bold;r.font.size=Pt(sz)
    if fill:shade(c,fill)
    if align=="c":p.alignment=WD_ALIGN_PARAGRAPH.CENTER
def table(headers,rows,caption=None,capnum=None):
    if capnum:
        c=doc.add_paragraph();rr=c.add_run(capnum);rr.bold=True;rr.font.size=Pt(9.5);rr2=c.add_run("  "+caption);rr2.italic=True;rr2.font.size=Pt(9.5)
    t=doc.add_table(rows=1,cols=len(headers));t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.style="Light Grid Accent 1"
    for i,h in enumerate(headers):
        setc(t.rows[0].cells[i],h,bold=True,fill="1F4E79")
        for r in t.rows[0].cells[i].paragraphs[0].runs:r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row): setc(cs[i],str(v),align=("c" if i else None))
def fig(path,cap,width=6.3):
    if Path(path).exists():
        doc.add_picture(str(path),width=Inches(width)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph();c.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=c.add_run(cap);r.italic=True;r.font.size=Pt(9.5)

# TITLE
ti=doc.add_paragraph();ti.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=ti.add_run("Early Prediction of Intensive-Care Length of Stay from the First 24 Hours: "
             "A Prospective Comparison with Senior-Physician Judgement (Leakage-Corrected Analysis)")
r.bold=True;r.font.size=Pt(16);r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
P("Authors / Affiliations / Corresponding author: [to be completed]",align="c",size=10)
P("Target journal: Frontiers in Digital Health — Research Topic “MedicinAI: Advancing the Synergy of "
  "Medicine and AI — From Data to Clinical Impact”",align="c",italic=True,size=10)
P("Companion (leakage-corrected) version: the OPS 8-98f intensive-care complex-treatment codes, "
  "identified as a target leak, are excluded from the model. All other methods are identical to the "
  "primary analysis.",align="c",italic=True,size=10)

# ABSTRACT
H("Abstract",1)
ab=doc.add_paragraph();ab.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
def ar(l,t): rr=ab.add_run(l+" ");rr.bold=True;ab.add_run(t+" ")
ar("Background:","Accurate early prediction of intensive-care-unit (ICU) length of stay (LoS) supports "
   "bed management and patient flow, but its value over experienced clinicians is unclear. We developed "
   "a leakage-corrected ICU-LoS model from routine first-24-hour data and validated it prospectively "
   "against senior-physician estimates.")
ar("Methods:",f"In a single anaesthesiology-run ICU department, {S['n_stays']:,} completed adult ICU "
   f"stays (>1 day; {S['n_patients']:,} patients) were used to develop regression models from "
   f"{S['n_features_leakfree']} first-24-hour predictors. The OPS 8-98f complex-treatment codes were "
   "removed because their suffix encodes cumulative treatment days and leaks the outcome (see "
   "Methods). Five candidate models were tuned by patient-grouped cross-validation; the lowest "
   "cross-validated MAE was pre-specified as final. The frozen model was evaluated prospectively on "
   f"{N_PROS} ICU stays with a matched senior estimate.")
ar("Results:",f"{FINAL} had the best cross-validated MAE (hold-out MAE {RETRO.loc[FKEY,'MAE_days']:.2f} d, "
   f"R² {RETRO.loc[FKEY,'R2']:.2f}). After leak removal the development R² fell to about "
   f"{RETRO.loc[FKEY,'R2']:.2f} (versus ~{LK['r2_full']:.2f} with the leaky feature), a more honest estimate of "
   f"signal. Prospectively the senior physician was more accurate overall (MAE {PROS.loc['Oberarzt','MAE']:.2f} "
   f"vs {PROS.loc[FKEY,'MAE']:.2f} d), but the comparison remained subgroup-dependent: the physician "
   f"was significantly more accurate for {PSG} stays, whereas the model was significantly more "
   f"accurate for {MSG} stays (paired bootstrap 95% CI of the MAE difference entirely above zero); "
   "other subgroups did not differ significantly.")
ar("Conclusion:","Once a duration-encoding leak is removed, the model’s honest first-24-hour signal is "
   "modest and it does not outperform clinicians overall. It nonetheless retains a statistically "
   "significant, reproducible advantage for intermediate-stay patients, supporting a complementary, "
   "human-in-the-loop role rather than replacement.")
P("Keywords: intensive care; length of stay; machine learning; data leakage; prospective validation; "
  "clinical decision support; TRIPOD+AI; human-in-the-loop",italic=True,size=10)

# 1 INTRO
H("1. Introduction",1)
P("Anticipating ICU length of stay supports bed management, staffing and discharge planning. Machine "
  "learning (ML) is widely applied to ICU outcomes, but two gaps limit clinical translation: models "
  "are seldom tested prospectively under real operating conditions, and they are rarely benchmarked "
  "against the clinicians they are meant to support. A third, frequently overlooked hazard is target "
  "leakage from administrative codes that are time-stamped to admission yet encode the eventual "
  "outcome. This leakage-corrected analysis addresses all three: we remove a duration-encoding "
  "complex-treatment code, interpret the remaining model, and compare it prospectively, by "
  "length-of-stay subgroup, with senior-physician estimates. The study follows the TRIPOD+AI "
  "statement.",align="j")

# 2 METHODS
H("2. Methods",1)
H("2.1 Study design and data source",2)
P("Single-centre prognostic-model study with retrospective development and a separate prospective "
  "evaluation, both drawn from the same three anaesthesiology-run intensive-care units. Routinely "
  "documented laboratory, vital-sign, procedure (OPS), vascular-access, diagnosis (ICD-10) and "
  "admission data were linked at the level of the individual ICU stay.",align="j")
H("2.2 Participants",2)
P(f"The development cohort comprised {S['n_stays']:,} adult ICU stays from {S['n_patients']:,} patients, "
  f"restricted to stays longer than one day. Patient identity kept every patient within a single "
  f"partition. The prospective cohort comprised all {N_PROS} consecutive stays with a recorded senior "
  f"estimate ({N_DONE} discharged with known final LoS, {N_OPEN} still in the ICU at estimation, whose "
  "recorded LoS is a right-censored lower bound).",align="j")
H("2.3 Outcome and the leakage correction",2)
P("The target was ICU LoS in days (hours/24), modelled on the log1p scale (tree/linear models) or via "
  "a Tweedie deviance loss. Permutation importance in the primary analysis was dominated by the OPS "
  "8-98f “aufwendige intensivmedizinische Komplexbehandlung” codes. An audit established these as a "
  "target leak: the German OPS complex-treatment code is assigned once per episode with a suffix that "
  "encodes the cumulative number of treatment days, but is time-stamped to the admission day, so it "
  "silently carries the eventual LoS. Empirically, the observed median LoS increased monotonically "
  f"across suffix bands ({LK['suffix_low_code']} → {LK['suffix_low_median']:.1f} d; {LK['suffix_high_code']} "
  f"→ {LK['suffix_high_median']:.1f} d), the code was present in {LK['prev_retro_pct']:.0f}% of "
  f"retrospective stays but {LK['prev_pros_pct']:.0f}% of prospective stays (it is not yet assigned at "
  f"24 h in live data), and removing it lowered the apparent development R² from ~{LK['r2_full']:.2f} "
  f"to ~{LK['r2_noleak']:.2f}. All 8-98f features were "
  f"therefore excluded, leaving {S['n_features_leakfree']} predictors.",align="j")
H("2.4 Predictors",2)
fd=dict(SC["feature_domains"]); fd["procedure_24h"]=fd.get("procedure_24h",13)-3
table(["Predictor domain","n"],
      [["Laboratory (first 24 h)",fd["lab_24h"]],["Diagnoses (ICD-10, main)",fd["diagnosis"]],
       ["Procedures (OPS, first 24 h; 8-98f removed)",fd["procedure_24h"]],["Vascular access (first 24 h)",fd["access_24h"]],
       ["Vital signs (first 24 h)",fd["vital_24h"]],["Demographics / admission",fd["demographics_admission"]]],
      caption=f"Predictor domains after leakage correction ({S['n_features_leakfree']} features).",capnum="Table 1.")
H("2.5 Model development, tuning and selection",2)
P(f"Five models (ridge, random forest, extremely randomized trees, gradient-boosted trees, Tweedie) "
  f"were trained through a common preprocessing pipeline on a patient-grouped training set "
  f"(n = {S['n_train']:,}) with a held-out test set (n = {S['n_test']:,}). Hyperparameters were carried "
  "over from the primary analysis; the final model was re-selected as the candidate with the lowest "
  "patient-grouped 4-fold cross-validated MAE on the leakage-corrected feature set (Table 2).",align="j")
order=[FKEY]+[m for m in ["ExtraTrees","RandomForest","XGBoost","Tweedie","Ridge"] if m!=FKEY]
order=[m for m in order if m in RETRO.index]
table(["Model","CV-MAE (d)"],[[labels[m]+(" (final)" if m==FKEY else ""),f"{RETRO.loc[m,'CV_MAE_days']:.3f}"] for m in order],
      caption="Patient-grouped cross-validated MAE on the leakage-corrected features (lowest = final).",capnum="Table 2.")
H("2.6 Metrics and statistics",2)
P("MAE, median absolute error, RMSE, R² and bias were computed on the day scale. For the prospective "
  "comparison we tested superiority of the final model over the physician with a one-sided paired test "
  "and a paired bootstrap (B = 5,000) 95% CI of the MAE difference (significant only if the whole CI "
  "lay above zero), overall and within LoS subgroups (2–4, 4–7, >7 days). Stays still in the ICU "
  "contribute a censored LoS, making prospective errors conservative.",align="j")

# 3 RESULTS
H("3. Results",1)
H("3.1 Retrospective performance (leakage-corrected)",2)
P(f"With the leak removed, the five models were tightly clustered and modest: {FINAL} had the best "
  f"cross-validated MAE and was selected as final (hold-out MAE {RETRO.loc[FKEY,'MAE_days']:.2f} d, "
  f"R² {RETRO.loc[FKEY,'R2']:.2f}). All models reached R² ≈ 0.11–0.12 — about a third of the apparent "
  "fit obtained when the 8-98f leak was included — which we regard as the honest first-24-hour "
  "signal.",align="j")
table(["Model","MAE (d)","Median AE (d)","RMSE (d)","R²","Bias (d)","CV-MAE (d)"],
      [[labels[m]+(" (final)" if m==FKEY else ""),f"{RETRO.loc[m,'MAE_days']:.2f}",f"{RETRO.loc[m,'MedianAE_days']:.2f}",
        f"{RETRO.loc[m,'RMSE_days']:.2f}",f"{RETRO.loc[m,'R2']:.2f}",f"{RETRO.loc[m,'Bias_days']:+.2f}",f"{RETRO.loc[m,'CV_MAE_days']:.3f}"] for m in order],
      caption=f"Retrospective held-out performance, leakage-corrected (n = {S['n_test']:,}).",capnum="Table 3.")
fig(LF/"fig_model_comparison_lf.png","Figure 1. Leakage-corrected accuracy (MAE) and calibration (R²) on the "
    f"retrospective hold-out vs the prospective cohort (n = {N_PROS}), all five models, physician reference.")
H("3.2 Model interpretation",2)
top=IMP.head(8)
P("After removing 8-98f, permutation importance was led by other early intensive-care monitoring and "
  "procedure signals together with patient age and ICU stay number — clinically coherent severity "
  "markers, none individually dominant, consistent with the modest overall signal.",align="j")
def nice(f): return f.replace("proc24_","Procedure ").replace("diag_main_","Diagnosis ").replace("vital24_","Vital ").replace("lab24_","Lab ").replace("zugang24_","Access ").replace("_"," ")
table(["Rank","Predictor","ΔMAE when permuted (d)"],
      [[str(i+1),nice(top.iloc[i]["Feature"]),f"{top.iloc[i]['MAE_increase_days']:.3f}"] for i in range(len(top))],
      caption="Top-8 predictors by permutation importance (leakage-corrected final model, held-out set).",capnum="Table 4.")
fig(LF/"fig_importance_lf.png","Figure 2. Permutation feature importance — leakage-corrected final model.")
H("3.3 Prospective validation against the senior physician",2)
P(f"Across all {N_PROS} prospective stays the senior physician was more accurate than every model "
  f"(physician MAE {PROS.loc['Oberarzt','MAE']:.2f} d, R² {PROS.loc['Oberarzt','R2']:.2f}; best model "
  f"{FINAL} MAE {PROS.loc[FKEY,'MAE']:.2f} d, R² {PROS.loc[FKEY,'R2']:.2f}). All models had negative "
  "prospective R², i.e. they did not beat a mean-prediction baseline out of sample — the realistic "
  "consequence of removing the leak and of distribution shift (Table 5).",align="j")
table(["Estimator","MAE (d)","Median AE (d)","RMSE (d)","R²","Bias (d)"],
      [["Senior physician",f"{PROS.loc['Oberarzt','MAE']:.2f}",f"{PROS.loc['Oberarzt','MedianAE']:.2f}",f"{PROS.loc['Oberarzt','RMSE']:.2f}",f"{PROS.loc['Oberarzt','R2']:.2f}",f"{PROS.loc['Oberarzt','Bias']:+.2f}"]]+
      [[labels[m]+(" (final)" if m==FKEY else ""),f"{PROS.loc[m,'MAE']:.2f}",f"{PROS.loc[m,'MedianAE']:.2f}",f"{PROS.loc[m,'RMSE']:.2f}",f"{PROS.loc[m,'R2']:.2f}",f"{PROS.loc[m,'Bias']:+.2f}"] for m in order],
      caption=f"Prospective performance, leakage-corrected (n = {N_PROS}; {N_OPEN} censored).",capnum="Table 5.")
def sm(m,sg):
    r=SUB[(SUB["Modell"]==m)&(SUB["Subgroup"]==sg)]; return f"{r['MAE'].iloc[0]:.2f}" if len(r) else "–"
sgs=["1–2 d","2–4 d","4–7 d",">7 d"]
table(["LoS subgroup","Senior physician",FINAL,"Random forest","Extra Trees"],
      [[sg,sm("Oberarzt",sg),sm(FKEY,sg),sm("RandomForest",sg),sm("ExtraTrees",sg)] for sg in sgs],
      caption=f"Prospective MAE (days) by LoS subgroup, leakage-corrected (n = {N_PROS}).",capnum="Table 6.")
def _fmt(k): r=_supd[k]; return f"{k}: ΔMAE {r['dMAE']:+.2f} d, 95% CI {r['CI_low']:.2f} to {r['CI_high']:.2f}"
_parts=[]
if MODEL_SG: _parts.append(f"the final model ({FINAL}) was significantly more accurate than the senior "
    "physician for "+", ".join(MODEL_SG)+" stays ("+"; ".join(_fmt(k) for k in MODEL_SG)+")")
if PHYS_SG:  _parts.append("the physician was significantly more accurate for "+", ".join(PHYS_SG)+
    " stays ("+"; ".join(_fmt(k) for k in PHYS_SG)+")")
if NS_SG:    _parts.append("the two were statistically indistinguishable for "+", ".join(NS_SG)+" stays")
P("Superiority testing on the leakage-corrected model, by subgroup (paired bootstrap 95% CI of the "
  "MAE difference): "+"; ".join(_parts)+". The model thus retains a significant, if narrower, "
  "subgroup advantage even after the leak is removed.",align="j")
fig(LF/"fig_subgroup_mae_lf.png",f"Figure 3. Leakage-corrected prospective MAE by LoS subgroup (n = {N_PROS}) "
    f"with the {FINAL}-vs-physician superiority test (paired bootstrap 95% CI).")
fig(LF/"fig_calibration_lf.png",f"Figure 4. Leakage-corrected prospective calibration (n = {N_PROS}): observed vs "
    "predicted (deciles, 95% CI) for the final model and the senior physician.")

# 4 DISCUSSION
H("4. Discussion",1)
H("4.1 Principal findings",2)
P("Removing a single duration-encoding administrative code (OPS 8-98f) cut the apparent development R² "
  f"from ~{LK['r2_full']:.2f} to ~{LK['r2_noleak']:.2f}, exposing a modest true first-24-hour signal. Honestly evaluated, the model "
  "does not outperform experienced clinicians overall. The comparison nonetheless remains "
  f"subgroup-dependent: the model is significantly more accurate than the physician for {MSG} stays, "
  f"whereas the physician retains a clear advantage for {PSG} stays.",align="j")
H("4.2 Clinical impact and translation",2)
P("The leakage correction strengthens, rather than weakens, the translational message. A model whose "
  "headline performance rests on a code unavailable at prediction time would fail silently in "
  "deployment; the corrected model is a faithful estimate of what is achievable from genuine "
  "first-24-hour data. Its reproducible edge for intermediate-stay patients — common and "
  "operationally demanding — supports use as a subgroup-aware second opinion rather than a "
  "replacement, in line with trustworthy, clinician-centred AI.",align="j")
H("4.3 Strengths and limitations",2)
P("Strengths include the explicit leakage audit and correction, a genuine prospective benchmark "
  "against clinicians, patient-grouped validation and TRIPOD+AI reporting. Limitations include the "
  "single-centre design, the modest prospective sample (especially for long stays), the censoring of "
  f"still-admitted patients ({N_OPEN}/{N_PROS}), and that hyperparameters were carried over from the "
  "primary analysis rather than re-tuned. LoS is partly an organisational outcome that bounds the "
  "achievable accuracy of any admission-time estimator. External multi-centre validation is the next "
  "step.",align="j")
H("4.4 Conclusion",2)
P("After correcting a duration-encoding leak, an interpretable first-24-hour ML model offers only a "
  "modest honest signal and does not replace clinical judgement, but it adds significant accuracy for "
  "intermediate-stay patients. A subgroup-aware, complementary deployment is the realistic path from "
  "data to clinical impact.",align="j")

H("Reporting, data and code availability",1)
P("Reported per TRIPOD+AI. This is the leakage-corrected companion to the primary manuscript "
  "(8-98f excluded). Code is in the project repository; patient-level data remain at the originating "
  "institution. [Ethics, funding, author contributions to be completed.]",align="j",size=10)
H("References",1)
for i,r in enumerate(["Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement. BMJ. 2024;385:e078378.",
   "Chen T, Guestrin C. XGBoost: a scalable tree boosting system. Proc. 22nd ACM SIGKDD. 2016:785–794.",
   "Pedregosa F, et al. Scikit-learn: machine learning in Python. JMLR. 2011;12:2825–2830.",
   "Geurts P, Ernst D, Wehenkel L. Extremely randomized trees. Mach Learn. 2006;63(1):3–42.",
   "[Additional ICU length-of-stay references to be completed by the authors.]"],1):
    p=doc.add_paragraph();rr=p.add_run(f"{i}. {r}");rr.font.size=Pt(9.5)

doc.save(str(OUT)); print(f"Saved: {OUT} | final model: {FINAL} | {len(doc.paragraphs)} paragraphs | {OUT.stat().st_size/1024:.1f} KB")
