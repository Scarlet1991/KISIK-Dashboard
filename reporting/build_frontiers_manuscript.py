# -*- coding: utf-8 -*-
"""
Erzeugt ein Manuskript (Original Research) im Frontiers-Stil als .docx
fuer den Research Topic 'MedicineAI: Advancing the Synergy of Medicine and AI'.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, RGBColor, Cm

OUT = Path(r"D:\Ausgangsdaten\KISIK Projekt\Eigene Auswertung\KISIK_Frontiers_Manuskript.docx")

FONT = "Times New Roman"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()

# Grundstil
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(6)

sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin = sec.bottom_margin = Cm(2.5)


# ----------------------------------------------------------------------
def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = ACCENT
    r.font.name = FONT


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = ACCENT
    r.font.name = FONT


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.italic = True
    r.font.size = Pt(12)
    r.font.name = FONT


def body(text, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = FONT
    return p


def labeled(label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(label + " ")
    r.bold = True
    r.font.name = FONT
    r2 = p.add_run(text)
    r2.font.name = FONT
    return p


def small(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = FONT
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


# ======================================================================
# TITELBLOCK
# ======================================================================
title("Predicting intensive-care length of stay from the first 24 hours: "
      "a leakage-controlled machine-learning study with prospective benchmarking "
      "against senior-physician judgement")

small("Manuscript prepared for the Frontiers Research Topic "
      "“MedicineAI: Advancing the Synergy of Medicine and AI – From Data to "
      "Clinical Impact.” Article type: Original Research. "
      "Author names, affiliations, corresponding author and ORCID identifiers to be inserted.")

# Autorenzeile (Platzhalter)
p = doc.add_paragraph()
r = p.add_run("Author One¹, Author Two¹, Author Three², Senior Author¹*")
r.font.name = FONT
p2 = doc.add_paragraph()
r = p2.add_run("¹ Department of Anaesthesiology and Intensive Care Medicine, [Institution], [City], [Country]\n"
               "² [Second affiliation], [City], [Country]\n"
               "* Correspondence: [name, e-mail address]")
r.font.size = Pt(11)
r.font.name = FONT


# ======================================================================
# ABSTRACT
# ======================================================================
h1("Abstract")

labeled("Introduction:",
        "Early and accurate prediction of intensive care unit (ICU) length of stay (LOS) could "
        "improve bed management, staffing and patient flow. Machine-learning (ML) models often "
        "report strong performance, yet two threats limit their clinical credibility: information "
        "leakage during model development, and the near-absence of head-to-head comparison with the "
        "experienced clinicians whose judgement they are intended to support.")

labeled("Methods:",
        "Using routine data from a single tertiary-care ICU clinical repository, we developed ML "
        "models to predict ICU LOS from information available within the first 24 hours after "
        "admission only. The retrospective development cohort comprised 17,032 ICU stays, split at "
        "the patient level (13,603 training / 3,429 test). Predictors (n = 102) covered demographics "
        "and admission context, first-24-hour laboratory values and vital signs, procedures, "
        "vascular-access devices and admission diagnoses. Four models – XGBoost, random forest, "
        "extra trees and ridge regression – were trained on a log-transformed target. Model "
        "predictions were then prospectively compared with prospectively documented senior-physician "
        "LOS estimates in an independent cohort, using 359 stays in which both were available; paired "
        "absolute errors were compared with Wilcoxon signed-rank tests.")

labeled("Results:",
        "Aggregating measurements across the whole admission produced optimistic apparent "
        "performance (coefficient of determination, R² ≈ 0.61); restricting predictors to a "
        "strict 24-hour window removed this leakage and revealed a substantially harder task. On the "
        "retrospective hold-out set, XGBoost performed best (mean absolute error [MAE] 2.11 days; "
        "R² 0.568). In the prospective comparison, however, the senior physician outperformed "
        "every ML model (MAE 2.60 vs 3.65 days for XGBoost; R² 0.251 vs −0.20) and produced "
        "the smaller absolute error in 70–74% of individual stays (all Wilcoxon p < 0.0001). "
        "Model errors were greatest for long stays (> 7 days), where predictions regressed toward the "
        "population mean.")

labeled("Conclusion:",
        "Under strict leakage control and prospective, clinician-benchmarked validation, 24-hour ML "
        "models did not match experienced senior-physician judgement for ICU LOS prediction. "
        "Translating ICU LOS models into clinical impact will require richer longitudinal inputs and "
        "routine benchmarking against clinicians, rather than against historical labels alone.")

p = doc.add_paragraph()
r = p.add_run("Keywords: ")
r.bold = True
r.font.name = FONT
r2 = p.add_run("intensive care unit; length of stay; machine learning; clinical prediction model; "
               "data leakage; prospective validation; clinical decision support; XGBoost")
r2.font.name = FONT


# ======================================================================
# 1 INTRODUCTION
# ======================================================================
h1("1  Introduction")

body("Intensive care is among the most resource-intensive components of hospital care, and the "
     "length of stay (LOS) on the intensive care unit (ICU) is a principal driver of cost, bed "
     "occupancy and staffing demand. Reliable early prediction of ICU LOS could therefore support "
     "operational decisions such as bed allocation, step-down planning and discharge coordination, "
     "and could help to set expectations for patients and relatives. Traditional severity-of-illness "
     "scores such as APACHE II, SAPS II and the SOFA score were developed primarily to predict "
     "mortality and to benchmark case mix, and they explain only a limited fraction of the variation "
     "in LOS [1–3].")

body("The growing availability of granular electronic health-record data and the maturation of "
     "machine-learning (ML) methods have prompted numerous attempts to predict ICU LOS from "
     "routinely collected data [4–6]. Reported performance is frequently encouraging. However, "
     "two recurring problems undermine the clinical credibility of such models. First, information "
     "leakage – the inadvertent inclusion of information that would not be available at the "
     "intended moment of prediction – can inflate apparent performance and is widespread in "
     "ML-for-health research [7]. For an outcome such as LOS, features aggregated over the entire "
     "admission are particularly hazardous, because their values are mechanically related to how long "
     "the patient ultimately stays. Second, prediction models are almost always evaluated against "
     "historical outcome labels, and only rarely against the experienced clinicians whose judgement "
     "they are meant to augment or replace. Without such a comparison it is impossible to know whether "
     "a model adds value at the bedside.")

body("Reporting and validation standards for clinical prediction models, including their AI-specific "
     "extensions, emphasise transparent handling of the prediction time point and rigorous external "
     "validation [8]. Prospective, clinician-benchmarked evaluation remains the exception rather than "
     "the rule.")

body("In this study we set out to develop a clinically realistic ML model for ICU LOS that uses only "
     "information available within the first 24 hours after admission, to quantify explicitly the "
     "effect of leakage from whole-stay feature aggregation, and – most importantly – to "
     "benchmark the resulting models prospectively against documented senior-physician estimates of "
     "LOS in an independent cohort. Our aim was not merely to report a performance metric, but to ask "
     "the question that matters for clinical impact: at the point of decision, does the model beat the "
     "doctor?")


# ======================================================================
# 2 MATERIALS AND METHODS
# ======================================================================
h1("2  Materials and Methods")

h2("2.1  Data source and setting")
body("We used de-identified routine-care data from the clinical data repository (KISIK) of a single "
     "tertiary-care centre. Several data modalities – ICU stay records, admission diagnoses "
     "(coded with the International Classification of Diseases, 10th revision; ICD-10), laboratory "
     "results, vital-sign measurements, procedure codes (the German Operationen- und "
     "Prozedurenschlüssel; OPS) and vascular-access devices – were linked by case "
     "identifier. The study was conducted in accordance with the Declaration of Helsinki; ethical "
     "approval and the waiver of informed consent for the analysis of de-identified routine data are "
     "to be inserted (see Ethics statement).")

h2("2.2  Cohort definition")
body("Eligible records were ICU stays on intensive care units of the participating departments "
     "(17 ward / care-unit combinations, predominantly anaesthesiological ICUs). To focus on stays "
     "for which LOS prediction is operationally meaningful, we restricted the analysis to stays "
     "longer than 24 hours (> 1 day). The retrospective development cohort comprised 17,032 ICU "
     "stays. To prevent optimistic bias from repeated admissions of the same patient, the data were "
     "split at the patient level (group-wise 80 / 20 split), yielding 13,603 stays for training and "
     "3,429 for hold-out testing, with no patient represented in both sets. An independent, "
     "later-period cohort of 5,655 ICU stays was reserved for prospective comparison "
     "(Section 2.6, Table 1).")

h2("2.3  Outcome")
body("The outcome was ICU LOS in days, modelled as a continuous variable. Because the distribution "
     "of LOS is strongly right-skewed, models were trained on a log-transformed target "
     "(log[1 + LOS]); predictions were back-transformed to the day scale for evaluation.")

h2("2.4  Feature engineering and leakage control")
body("All predictors were derived strictly from the first 24 hours after ICU admission (from the "
     "admission timestamp to admission + 24 h). This design choice reflects the clinical setting in "
     "which an early estimate is required, and it is the central safeguard against leakage: no "
     "measurement, procedure or device recorded after the prediction window can enter the model. "
     "The final feature set comprised 102 predictors across six domains: demographics and admission "
     "context (age, ICU stay number, admission hour / weekday / month, care-unit type); first-24-hour "
     "laboratory values (e.g. haemoglobin, bilirubin, potassium, sodium, lactate, glucose, summarised "
     "as first, mean, minimum, maximum, last and count); first-24-hour vital signs (peripheral oxygen "
     "saturation); procedures (OPS, binary presence and count); vascular-access devices (arterial "
     "line, central venous catheter, urinary catheter; binary presence and count); and admission "
     "diagnoses (ICD-10; binary). The composition of the feature set is summarised in Table 1.")
body("To quantify the impact of leakage, we contrasted this leakage-controlled feature set with a "
     "naive alternative in which the same measurements were aggregated over the entire admission "
     "rather than over the first 24 hours. Whole-stay summaries (for example the last or maximum "
     "value of a laboratory parameter, or the total number of vascular-access devices) are "
     "mechanically correlated with LOS and thus constitute leakage when used to predict it.")

h2("2.5  Model development")
body("Four regression models spanning a range of complexity were trained: extreme gradient boosting "
     "(XGBoost) [9], random forest, extremely randomised trees (extra trees) and ridge regression. "
     "All models shared an identical preprocessing pipeline (median imputation of missing numeric "
     "values, standardisation where appropriate, and one-hot encoding of categorical variables) and "
     "the same log-transformed target. Hyperparameters were kept at sensible, lightly tuned defaults "
     "to avoid over-fitting to the development data. Models were fitted on the training set and "
     "evaluated once on the held-out test set.")

h2("2.6  Prospective comparison with senior-physician estimates")
body("In the independent prospective cohort, senior physicians had documented an expected ICU LOS "
     "for individual patients as part of routine clinical workflow (360 estimates). Each model was "
     "applied to the prospective stays using the identical 24-hour feature definition, and model "
     "predictions were matched to the documented senior-physician estimate on the stay identifier. "
     "After matching and removal of records with missing values, 359 stays remained in which both a "
     "model prediction and a senior-physician estimate were available; these formed the matched "
     "evaluation set for the head-to-head comparison.")

h2("2.7  Statistical analysis")
body("Predictive performance was summarised by the MAE, the median absolute error, the root-mean-"
     "square error (RMSE), the coefficient of determination (R²) and the mean signed bias "
     "(mean of predicted minus observed). For the prospective head-to-head comparison, the paired "
     "absolute errors of each model and of the senior physician were compared with the Wilcoxon "
     "signed-rank test; in addition, we report the proportion of individual stays in which the "
     "physician produced the smaller absolute error. Performance was further examined within strata "
     "of observed LOS (1–7, > 7 and > 14 days) and by means of a calibration analysis comparing "
     "the mean predicted with the mean observed LOS across LOS strata. Analyses were performed in "
     "Python (scikit-learn, XGBoost, SciPy). A two-sided p-value < 0.05 was considered statistically "
     "significant.")


# ======================================================================
# 3 RESULTS
# ======================================================================
h1("3  Results")

h2("3.1  Cohort")
body("The retrospective development cohort comprised 17,032 ICU stays (13,603 training / 3,429 "
     "hold-out test). The prospective cohort comprised 5,655 stays, of which 359 could be matched to "
     "a documented senior-physician LOS estimate and formed the matched evaluation set. Cohort "
     "definitions, data sources and the predictor domains are summarised in Table 1.")

h2("3.2  Impact of leakage control")
body("The choice of feature window had a marked effect on apparent performance. A feature set that "
     "summarised measurements across the whole admission produced optimistic estimates "
     "(R² ≈ 0.61). Restricting predictors to a strict first-24-hour window – so that no "
     "post-window information could enter the model – substantially reduced apparent performance "
     "and exposed the genuine difficulty of predicting LOS from early data. All results reported "
     "hereafter use the leakage-controlled 24-hour feature set.")

h2("3.3  Retrospective model performance")
body("On the patient-level hold-out test set (n = 3,429), XGBoost was the best-performing model "
     "(MAE 2.11 days, median absolute error 0.92 days, RMSE 4.22 days, R² 0.568). Ridge "
     "regression and random forest performed similarly to one another (R² 0.370 and 0.360, "
     "MAE 2.56 and 2.55 days), and extra trees ranked last (R² 0.298, MAE 2.70 days). All "
     "models showed a modest negative bias, tending to under-estimate LOS. Full results are given in "
     "Table 2 and Figure 1.")

h2("3.4  Prospective comparison with senior physicians")
body("In the matched prospective cohort (n = 359), the senior physician outperformed every ML model "
     "on all overall metrics (MAE 2.60 days, median absolute error 0.93 days, R² 0.251). The "
     "best model, XGBoost, achieved an MAE of 3.65 days and a negative R² (−0.20), "
     "indicating poorer agreement than simply predicting the cohort mean; the remaining models "
     "performed similarly or worse (Table 3A). The physician produced the smaller absolute error in "
     "70–74% of individual stays, depending on the comparator model, and the difference in paired "
     "absolute errors was statistically significant for every model (all Wilcoxon p < 0.0001; "
     "Table 3A, C). The notable drop in model performance from the retrospective hold-out to the "
     "prospective cohort is consistent with distribution shift between the development and prospective "
     "periods (Figure 1 versus Figures 2–4).")

h2("3.5  Performance across length-of-stay strata and calibration")
body("Stratified analysis showed that the physician’s advantage was driven by long stays. For "
     "stays of 1–7 days (n = 216) the senior physician and the models were close (MAE 1.39 days "
     "for the physician versus 1.51–1.61 days for the tree-based models). For stays longer than "
     "7 days (n = 70) and longer than 14 days (n = 30) the gap widened sharply (physician MAE 7.74 "
     "and 11.77 days versus 11.6–12.2 and 17.2–18.9 days for the models; Table 3B, "
     "Figure 2). The calibration analysis (Figure 5) made the mechanism explicit: across LOS strata "
     "the mean model prediction remained almost flat (around 2.5–3.3 days regardless of the true "
     "stratum), whereas the mean physician estimate tracked the observed mean far more closely. The "
     "models therefore regressed long-stay patients toward the population mean, while the physicians "
     "retained discriminative information about which patients would stay long. The Bland–Altman "
     "analysis (Figure 3) confirmed a larger and more LOS-dependent bias for the models than for the "
     "physician.")


# ======================================================================
# 4 DISCUSSION
# ======================================================================
h1("4  Discussion")

body("In a large single-centre cohort we developed ML models to predict ICU LOS from the first 24 "
     "hours after admission, explicitly controlled for information leakage, and benchmarked the "
     "models prospectively against documented senior-physician estimates. Three findings stand out. "
     "First, leakage from whole-stay feature aggregation materially inflated apparent performance, "
     "and removing it revealed a much harder task. Second, on a leakage-controlled retrospective "
     "hold-out set, gradient boosting (XGBoost) achieved the best performance. Third, and most "
     "importantly, in a prospective head-to-head comparison the experienced senior physician "
     "outperformed every ML model, and did so significantly, with the advantage concentrated among "
     "long-staying patients.")

body("The leakage result is a cautionary tale that generalises well beyond ICU LOS. Because LOS is a "
     "temporal outcome, any feature summarising the whole admission encodes part of the answer. Such "
     "features are easy to construct inadvertently from wide, pre-aggregated data tables, and they "
     "yield impressive but meaningless metrics. Our experience supports recent warnings that leakage "
     "is a leading cause of over-optimism and irreproducibility in ML-for-health research [7], and it "
     "underlines why the prediction time point must be defined and enforced before any feature is "
     "computed [8].")

body("The central message of our study concerns the comparison with clinicians. Most LOS models are "
     "validated only against historical labels; under that standard, our XGBoost model would appear "
     "reasonable. Benchmarked against the people who actually make the decision, however, it fell "
     "short. The calibration analysis suggests why: with only 24 hours of data, the models could not "
     "identify which patients would become long-stayers and defaulted toward the mean, whereas "
     "experienced physicians integrate contextual knowledge – surgical trajectory, anticipated "
     "complications, social and organisational factors – that is not captured in structured "
     "early data. For short, uncomplicated stays the model and the clinician were comparable; the "
     "clinical value of an early predictor lies precisely in the difficult long-stay cases, and there "
     "the models were weakest.")

body("These findings have practical implications for the synergy of medicine and AI. A model that "
     "does not beat the clinician on the cases that matter is unlikely to deliver clinical impact if "
     "deployed as an autonomous predictor, and may erode trust. The more promising path is "
     "augmentation: surfacing the structured-data signal where it is reliable (short-stay triage, "
     "early flags) while preserving clinician judgement for complex cases, and routinely benchmarking "
     "any deployed model against clinicians rather than against historical labels alone. Improving the "
     "models will likely require richer longitudinal inputs – dynamic, repeatedly updated "
     "predictions that ingest data beyond the first 24 hours – and dedicated modelling of the "
     "long-stay tail.")

h2("4.1  Limitations")
body("This study has several limitations. It is based on data from a single centre, and the absolute "
     "performance figures may not transfer to other settings. The prospective comparison rested on "
     "359 matched stays with documented senior-physician estimates; although sufficient to "
     "demonstrate a significant and consistent difference, this is a modest sample, and the "
     "physicians’ estimates may themselves have been informed by evolving information rather than "
     "fixed at 24 hours. The marked performance drop between the retrospective and prospective "
     "cohorts indicates distribution shift, the precise drivers of which (case mix, coding practice, "
     "time period) were not fully characterised here. We deliberately restricted predictors to the "
     "first 24 hours and used lightly tuned models; more extensive feature engineering, longitudinal "
     "modelling or hyperparameter optimisation might narrow the gap, and represents an important "
     "direction for future work. Finally, we evaluated LOS as a point prediction; probabilistic or "
     "time-to-event formulations may be better suited to the long-stay tail.")


# ======================================================================
# 5 CONCLUSION
# ======================================================================
h1("5  Conclusion")
body("Under strict leakage control and prospective, clinician-benchmarked validation, machine-"
     "learning models built from the first 24 hours of ICU data did not match the judgement of "
     "experienced senior physicians for length-of-stay prediction, with the clinician’s "
     "advantage concentrated in the long-stay cases that matter most operationally. Realising the "
     "clinical impact promised by AI in critical care will depend less on incremental gains against "
     "historical labels and more on rigorous validation, transparent handling of the prediction time "
     "point, and honest, routine comparison with the clinicians the technology is meant to serve.")


# ======================================================================
# STATEMENTS
# ======================================================================
h1("Statements")

labeled("Data availability statement.",
        "The data analysed in this study are de-identified routine clinical data and are not publicly "
        "available owing to data-protection regulations. Requests to access the data may be directed "
        "to the corresponding author and are subject to institutional and regulatory approval.")

labeled("Ethics statement.",
        "The study involved the analysis of de-identified routine clinical data. The protocol and the "
        "waiver of informed consent were reviewed by [name of ethics committee / institutional review "
        "board], approval number [to be inserted]. The study was conducted in accordance with the "
        "Declaration of Helsinki.")

labeled("Author contributions.",
        "[To be completed, e.g.] Conceptualisation: [..]; data curation and analysis: [..]; "
        "methodology and software: [..]; clinical interpretation: [..]; writing – original draft: "
        "[..]; writing – review and editing: all authors. All authors read and approved the final "
        "manuscript.")

labeled("Funding.",
        "[To be completed. If none: “The authors declare that no funding was received for this "
        "work.”]")

labeled("Conflict of interest.",
        "The authors declare that the research was conducted in the absence of any commercial or "
        "financial relationships that could be construed as a potential conflict of interest. "
        "[Amend as appropriate.]")

labeled("Acknowledgements.",
        "[Optional – to be completed.]")


# ======================================================================
# REFERENCES  (echte Anker-Referenzen; bibliografische Details pruefen)
# ======================================================================
h1("References")

refs = [
    "Knaus WA, Draper EA, Wagner DP, Zimmerman JE. APACHE II: a severity of disease classification "
    "system. Crit Care Med. 1985;13(10):818–829.",
    "Le Gall JR, Lemeshow S, Saulnier F. A new Simplified Acute Physiology Score (SAPS II) based on "
    "a European/North American multicenter study. JAMA. 1993;270(24):2957–2963.",
    "Vincent JL, Moreno R, Takala J, et al. The SOFA (Sepsis-related Organ Failure Assessment) "
    "score to describe organ dysfunction/failure. Intensive Care Med. 1996;22(7):707–710.",
    "Verburg IWM, Atashi A, Eslami S, et al. Which models can I use to predict adult ICU length of "
    "stay? A systematic review. Crit Care Med. 2017;45(2):e222–e231.",
    "Johnson AEW, Pollard TJ, Shen L, et al. MIMIC-III, a freely accessible critical care database. "
    "Sci Data. 2016;3:160035.",
    "Pollard TJ, Johnson AEW, Raffa JD, et al. The eICU Collaborative Research Database, a freely "
    "available multi-center database for critical care research. Sci Data. 2018;5:180178.",
    "Kapoor S, Narayanan A. Leakage and the reproducibility crisis in machine-learning-based "
    "science. Patterns. 2023;4(9):100804.",
    "Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent reporting of a multivariable "
    "prediction model for individual prognosis or diagnosis (TRIPOD): the TRIPOD statement. "
    "Ann Intern Med. 2015;162(1):55–63.",
    "Chen T, Guestrin C. XGBoost: a scalable tree boosting system. In: Proceedings of the 22nd ACM "
    "SIGKDD International Conference on Knowledge Discovery and Data Mining (KDD ’16); 2016. "
    "p. 785–794.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"[{i}]  {ref}")
    r.font.size = Pt(10.5)
    r.font.name = FONT

small("Note: the reference list above contains landmark works as a starting point. Bibliographic "
      "details (volume, pages, year) should be verified, and topic-specific references relevant to "
      "your discussion should be added before submission. In-text markers [1]–[9] correspond to "
      "this list.")


# ----------------------------------------------------------------------
doc.save(str(OUT))
print(f"Gespeichert: {OUT}")
print(f"Groesse: {round(OUT.stat().st_size/1024, 1)} KB")
